from __future__ import annotations

import argparse
import io
import os
import time
import uuid
import json
from datetime import datetime
from pathlib import Path
from threading import Thread
from typing import Optional, List, Dict, Any
import re
import logging
from dotenv import load_dotenv

import numpy as np
import requests
import torch
from flask import Flask, jsonify, request, send_from_directory, send_file, url_for
from openai import OpenAI
from PIL import Image
from transformers import AutoModel, AutoImageProcessor
from werkzeug.utils import secure_filename
from groq import Groq
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import psycopg2
from psycopg2.extras import RealDictCursor, Json
from pgvector.psycopg2 import register_vector
from flask import url_for

load_dotenv()
logging.basicConfig(level=logging.INFO)
DB_DSN = "postgresql://administrationSTS:St%24%400987@avo-adb-002.postgres.database.azure.com:5432/Micrographie_IA"

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_BASE_DIR = BASE_DIR / "embeddings_v7"
IMAGES_DIR = OUTPUT_BASE_DIR / "images"
TEMP_UPLOAD_DIR = BASE_DIR / "temp_uploads"
DOCX_TEMP_DIR = BASE_DIR / "temp_docx"

OUTPUT_BASE_DIR.mkdir(parents=True, exist_ok=True)
IMAGES_DIR.mkdir(parents=True, exist_ok=True)
TEMP_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
DOCX_TEMP_DIR.mkdir(parents=True, exist_ok=True)

HARDCODED_OPENAI_API_KEY = ""
openai_api_key = HARDCODED_OPENAI_API_KEY or os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=openai_api_key) if openai_api_key else None

GROQ_API_KEYS = [
    "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",
    "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",
]

current_groq_key_index = 0
groq_api_key = GROQ_API_KEYS[current_groq_key_index] if GROQ_API_KEYS else os.getenv("GROQ_API_KEY")
groq_client = Groq(api_key=groq_api_key) if groq_api_key else None


def rotate_groq_key():
    global current_groq_key_index, groq_client, groq_api_key
    current_groq_key_index = (current_groq_key_index + 1) % len(GROQ_API_KEYS)
    new_key = GROQ_API_KEYS[current_groq_key_index]
    print(f"🔄 Rotation vers la clé Groq #{current_groq_key_index + 1}")
    groq_api_key = new_key
    groq_client = Groq(api_key=new_key)
    return groq_client


def call_groq_with_retry(messages, model="llama-3.3-70b-versatile", temperature=0.3, max_tokens=8000, response_format=None):
    if not groq_client:
        raise Exception("Groq client not initialized")
    attempts = 0
    max_attempts = len(GROQ_API_KEYS)
    while attempts < max_attempts:
        try:
            kwargs = {"model": model, "messages": messages, "temperature": temperature, "max_tokens": max_tokens}
            if response_format:
                kwargs["response_format"] = response_format
            response = groq_client.chat.completions.create(**kwargs)
            return response
        except Exception as e:
            error_message = str(e).lower()
            if "authentication" in error_message or "invalid" in error_message or "unauthorized" in error_message or "401" in error_message:
                print(f"⚠️ Erreur d'authentification avec la clé #{current_groq_key_index + 1}: {e}")
                attempts += 1
                if attempts < max_attempts:
                    rotate_groq_key()
                else:
                    raise Exception(f"❌ Toutes les clés Groq ({max_attempts}) ont échoué. Dernière erreur: {e}")
            else:
                raise e
    raise Exception("Échec de l'appel à Groq après rotation de toutes les clés")


DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
DINO_MODEL_NAME = "facebook/dinov2-large"
DINO_MODEL: Optional[AutoModel] = None
DINO_PROCESSOR: Optional[AutoImageProcessor] = None


def ensure_dino_loaded():
    global DINO_MODEL, DINO_PROCESSOR
    if DINO_MODEL is not None and DINO_PROCESSOR is not None:
        return
    print(f"🔧 Loading DINOv2 on {DEVICE}...")
    DINO_MODEL = AutoModel.from_pretrained(DINO_MODEL_NAME).to(DEVICE).eval()
    DINO_PROCESSOR = AutoImageProcessor.from_pretrained(DINO_MODEL_NAME)
    print("✅ DINOv2 loaded")


def compute_embedding_from_pil(image: Image.Image) -> np.ndarray:
    ensure_dino_loaded()
    image = image.convert("RGB")
    inputs = DINO_PROCESSOR(images=image, return_tensors="pt")
    inputs = {k: v.to(DEVICE) for k, v in inputs.items()}
    with torch.no_grad():
        outputs = DINO_MODEL(**inputs)
        embedding = outputs.last_hidden_state[:, 0, :].squeeze().cpu().numpy()
    return embedding.astype("float32")


ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg"}


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def guess_extension_from_mime(mime_type: Optional[str]) -> Optional[str]:
    if not mime_type:
        return None
    mt = mime_type.lower()
    if "png" in mt:
        return ".png"
    if "jpeg" in mt or "jpg" in mt:
        return ".jpg"
    return None


def cleanup_old_files(interval: int = 1800, max_age_seconds: int = 2 * 3600):
    while True:
        now = time.time()
        try:
            for f in TEMP_UPLOAD_DIR.iterdir():
                if not f.is_file():
                    continue
                try:
                    if now - f.stat().st_mtime > max_age_seconds:
                        f.unlink(missing_ok=True)
                except Exception as e:
                    print(f"Error deleting {f.name}: {e}")
            for f in DOCX_TEMP_DIR.iterdir():
                if not f.is_file() or not f.suffix == ".docx":
                    continue
                try:
                    if now - f.stat().st_mtime > 3600:
                        f.unlink(missing_ok=True)
                except Exception as e:
                    print(f"Error deleting {f.name}: {e}")
        except Exception as e:
            print(f"Cleanup error: {e}")
        time.sleep(interval)


cleanup_thread = Thread(target=cleanup_old_files, daemon=True)
cleanup_thread.start()


def serialize_to_json_compatible(obj):
    from datetime import datetime, date, time
    if isinstance(obj, dict):
        return {k: serialize_to_json_compatible(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [serialize_to_json_compatible(item) for item in obj]
    elif isinstance(obj, (datetime, date, time)):
        return obj.isoformat()
    else:
        return obj


def get_db_conn():
    conn = psycopg2.connect(DB_DSN)
    register_vector(conn)
    return conn


def search_similar_in_db(query_embedding: np.ndarray, top_k: int = 5) -> List[Dict[str, Any]]:
    query_vec = query_embedding.tolist()
    sql = """
        SELECT mi.id, mi.image_path, mi.matiere_id, m.nom_matiere, m.reference,
               (1 - (mi.embedding <=> %s)) AS similarity
        FROM public.matiere_images mi
        JOIN public.matieres m ON m.matiere_id = mi.matiere_id
        ORDER BY mi.embedding <=> %s
        LIMIT %s;
    """
    conn = get_db_conn()
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute(sql, (query_vec, query_vec, top_k))
            return [dict(r) for r in cur.fetchall()]
    finally:
        conn.close()


def build_image_url(image_path: str) -> str:
    filename = Path(image_path).name
    url = f"{request.host_url.rstrip('/')}/images/{secure_filename(filename)}"
    if url.startswith("http://"):
        url = "https://" + url[len("http://"):]
    return url


def generate_fiche_adn_content_with_groq(fiche_data, material_name, reference, type_matiere, specifications):
    if not groq_client:
        return generate_fallback_fiche_adn_content(material_name, reference, type_matiere, specifications)
    try:
        datasheet_spec = msds_spec = lab_control_spec = None
        expert_notes_data = []
        fiches_data = []
        if isinstance(specifications, dict):
            specs_list = specifications.get("specifications", [])
            raw_expert_notes = specifications.get("expert_notes", [])
            fiches_data = specifications.get("fiches", [])
            for note in raw_expert_notes:
                if isinstance(note, dict):
                    note_json = note.get("note_json", {})
                    if note_json:
                        expert_notes_data.append({"expert_notes": note_json.get("expert_notes", ""), "full_text": note_json.get("full_text", ""), "magnification": note_json.get("magnification", ""), "protocol": note_json.get("protocol", "")})
        else:
            specs_list = specifications if isinstance(specifications, list) else []
        for spec in specs_list:
            if isinstance(spec, dict):
                source_type = spec.get("source_type", "").lower()
                if "datasheet" in source_type:
                    datasheet_spec = spec
                elif "msds" in source_type:
                    msds_spec = spec
                elif "control" in source_type or "feuille" in source_type:
                    lab_control_spec = spec
        prompt_data = {"material": {"nom_matiere": material_name, "reference": reference, "type_matiere": type_matiere}, "datasheet": datasheet_spec.get("donnees", {}) if datasheet_spec else {}, "msds": msds_spec.get("donnees", {}) if msds_spec else {}, "lab_control": lab_control_spec.get("donnees", {}) if lab_control_spec else {}, "expert_notes": expert_notes_data[:5] if expert_notes_data else [], "fiches": fiches_data}
        prompt = f"""Generate a COMPLETE and PROFESSIONAL MATERIAL DNA SHEET (FICHE ADN) in ENGLISH with the following strict structure.
Use ONLY the data provided in JSON. If a section lacks data, write "Not available".
⚠️ MANDATORY TRANSLATION RULE: ANY TEXT OR FIELD taken from the input data MUST be translated to ENGLISH
COMPLETE JSON DATA: {json.dumps(prompt_data, ensure_ascii=False, default=str, indent=2)}
STRICT FORMATTING REQUIREMENTS:
I — IDENTITY & LOGISTICS
II — GENERAL PRODUCT CHARACTERISTICS
III — CHEMICAL PROPERTIES - TRIPARTITE STRUCTURE
### III.1 QUANTIFIED PROPERTIES (Datasheet)
### III.2 DETAILED COMPOSITION (MSDS §2)
### III.3 STABILITY & HAZARDS (MSDS §9, §10, §3)
## IV — PHYSICAL PROPERTIES
## V — LASER GRANULOMETRY
## VI — GRANULOMETRIC CONTROLS (LAB-CONTROL)
## VII — EXPERT NOTES & OBSERVATIONS
## VIII — STORAGE
## IX — PACKAGING
## X — SAFETY — COMPLETE MSDS DATA
OUTPUT FORMAT: Use Markdown with tables, bullet lists, **bold** for keywords. ALL TEXT IN ENGLISH ONLY."""
        message = call_groq_with_retry(messages=[{"role": "user", "content": prompt}], model="llama-3.3-70b-versatile", max_tokens=5000)
        content = message.choices[0].message.content if message.choices else ""
        if not content or len(content.strip()) < 100:
            raise Exception("Groq returned insufficient content")
        return content
    except Exception as e:
        print(f"❌ Groq generation failed: {e}")
        raise e


def add_formatted_markdown_to_docx(doc: Document, markdown_text):
    lines = markdown_text.strip().replace('\r\n', '\n').split('\n')
    in_table = False
    table = None
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                header_line = line
                if i + 1 < len(lines) and lines[i+1].strip().startswith('|--'):
                    num_cols = len([h.strip() for h in header_line.split('|') if h.strip()])
                    if num_cols > 0:
                        table = doc.add_table(rows=1, cols=num_cols)
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        headers = [h.strip() for h in header_line.split('|') if h.strip()]
                        for j, header in enumerate(headers):
                            if j < num_cols:
                                hdr_cells[j].text = header
                        i += 2
                        continue
                else:
                    in_table = False
            else:
                if line.startswith('|'):
                    row_data = [cell.strip() for cell in line.split('|') if cell.strip()]
                    if table and len(row_data) == table.columns:
                        row_cells = table.add_row().cells
                        for j, cell_text in enumerate(row_data):
                            row_cells[j].text = cell_text
                        i += 1
                        continue
                    else:
                        in_table = False
                        table = None
                else:
                    in_table = False
                    table = None
        if not in_table:
            if line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith(('* ', '- ')):
                indent_level = (len(line) - len(line.lstrip(' '))) // 2
                style = 'List Bullet'
                if indent_level > 0:
                    style = f'List Bullet {indent_level + 1}'
                doc.add_paragraph(line[2:].strip(), style=style)
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        p.add_run(part[2:-2]).bold = True
                    elif part:
                        p.add_run(part)
        i += 1


def generate_fallback_fiche_adn_content(material_name, reference, type_matiere, specifications):
    content = f"FICHE ADN - MATIÈRE\n\nNom: {material_name}\nRéférence: {reference}\nType: {type_matiere}\n\nSPÉCIFICATIONS TECHNIQUES\n\n"
    if isinstance(specifications, list):
        for spec in specifications:
            if isinstance(spec, dict):
                content += f"• {spec.get('source_type', 'Donnée')}: {spec.get('donnees', 'N/A')}\n"
    content += "\n\nRECOMMANDATIONS D'UTILISATION\n\nCette matière doit être manipulée selon les spécifications techniques ci-dessus.\n"
    return content


def get_first_image_for_material(matiere_id: int) -> Optional[Image.Image]:
    conn = None
    try:
        conn = get_db_conn()
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT image_path FROM public.matiere_images WHERE matiere_id = %s LIMIT 1", (matiere_id,))
            result = cur.fetchone()
            if result and result.get("image_path"):
                image_path = result["image_path"]
                normalized_path = image_path.replace("\\", "/")
                filename = Path(normalized_path).name
                for file_path in [Path(normalized_path), BASE_DIR / normalized_path, IMAGES_DIR / filename, BASE_DIR / "output_v3" / "images" / filename]:
                    if file_path.exists():
                        try:
                            return Image.open(file_path).convert("RGB")
                        except Exception:
                            pass
        return None
    except Exception as e:
        print(f"⚠️ Error retrieving image: {e}")
        return None
    finally:
        if conn:
            conn.close()


def get_all_images_for_material(matiere_id: int, limit: int = 2) -> List[Dict[str, Any]]:
    conn = None
    try:
        conn = get_db_conn()
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT mi.id, mi.image_path, men.note_json
                FROM public.matiere_images mi
                LEFT JOIN public.matiere_expert_notes men ON men.matiere_image_id = mi.id
                WHERE mi.matiere_id = %s ORDER BY mi.id LIMIT %s
            """, (matiere_id, limit))
            results = cur.fetchall()
            images_data = []
            for result in results:
                image_path = result.get("image_path")
                note_json = result.get("note_json") or {}
                magnification = note_json.get("magnification", "N/A") if isinstance(note_json, dict) else "N/A"
                image_obj = None
                if image_path:
                    normalized_path = image_path.replace("\\", "/")
                    filename = Path(normalized_path).name
                    for file_path in [Path(normalized_path), BASE_DIR / normalized_path, IMAGES_DIR / filename, BASE_DIR / "output_v3" / "images" / filename]:
                        try:
                            if file_path.exists():
                                image_obj = Image.open(file_path).convert("RGB")
                                break
                        except Exception:
                            pass
                images_data.append({"image_path": image_path, "magnification": magnification, "image_obj": image_obj})
            return images_data
    except Exception as e:
        print(f"⚠️ Error retrieving images: {e}")
        return []
    finally:
        if conn:
            conn.close()


# =============================================================================
# ROOT / HEALTH
# =============================================================================

@app.route("/", methods=["GET"])
def root():
    return jsonify({"service": "micrograph-search-api", "status": "ok", "model": DINO_MODEL_NAME, "dino_loaded": DINO_MODEL is not None, "images_dir": str(IMAGES_DIR)}), 200


@app.route("/health", methods=["GET"])
def health():
    check_db = request.args.get("check_db", "false").strip().lower() in {"1", "true", "yes"}
    db_ok = db_error = None
    if check_db:
        try:
            conn = get_db_conn()
            with conn.cursor() as cur:
                cur.execute("SELECT 1")
                cur.fetchone()
            conn.close()
            db_ok = True
        except Exception as e:
            db_ok = False
            db_error = str(e)
    return jsonify({"status": "ok", "dino_loaded": DINO_MODEL is not None, "groq_configured": groq_client is not None, "openai_configured": client is not None, "db_ok": db_ok, "db_error": db_error}), 200


@app.route("/images/<path:filename>", methods=["GET"])
def serve_image(filename):
    try:
        return send_from_directory(str(IMAGES_DIR), filename)
    except Exception:
        return jsonify({"error": "not_found"}), 404


@app.route("/temp_files/<path:filename>", methods=["GET"])
def serve_temp_file(filename):
    try:
        return send_from_directory(str(TEMP_UPLOAD_DIR), filename)
    except Exception:
        return jsonify({"error": "temp_file_not_found"}), 404


# =============================================================================
# UPLOAD AND SEARCH (MATIERES)
# =============================================================================

@app.route("/upload_and_search", methods=["POST"])
def upload_and_search():
    data = request.get_json(silent=True) or {}
    refs = data.get("openaiFileIdRefs")
    top_k = int(data.get("top_k", 5))
    if not refs or not isinstance(refs, list):
        return jsonify({"success": False, "error": "missing_openaiFileIdRefs"}), 400
    if top_k < 1 or top_k > 50:
        return jsonify({"success": False, "error": "invalid_top_k"}), 400
    final_results = []
    errors = []
    for file_ref in refs:
        try:
            if not isinstance(file_ref, dict):
                errors.append("Each item must be an object.")
                continue
            file_id = file_ref.get("id")
            download_link = file_ref.get("download_link")
            original_name = file_ref.get("name") or "uploaded_file"
            mime_type = file_ref.get("mime_type")
            if not file_id:
                errors.append("Missing id in file reference.")
                continue
            file_bytes = None
            if download_link:
                try:
                    r = requests.get(download_link, timeout=20)
                    r.raise_for_status()
                    file_bytes = r.content
                except Exception as e:
                    print(f"⚠️ download_link failed: {e}")
            if file_bytes is None:
                if not client:
                    errors.append(f"{original_name}: OpenAI API key not configured.")
                    continue
                file_bytes = client.files.content(file_id).read()
            filename_safe = secure_filename(original_name or "uploaded_file") or "uploaded_file"
            if "." not in filename_safe:
                ext = guess_extension_from_mime(mime_type) or ".png"
                filename_safe += ext
            if not allowed_file(filename_safe):
                errors.append(f"{original_name}: File type not allowed.")
                continue
            unique_filename = f"{uuid.uuid4().hex}_{int(time.time())}_{filename_safe}"
            file_path = TEMP_UPLOAD_DIR / unique_filename
            with open(file_path, "wb") as f:
                f.write(file_bytes)
            file_url = f"{request.host_url.rstrip('/')}/temp_files/{unique_filename}"
            if file_url.startswith("http://"):
                file_url = "https://" + file_url[len("http://"):]
            img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
            query_embedding = compute_embedding_from_pil(img)
            rows = search_similar_in_db(query_embedding, top_k=top_k)
            search_results = [{"id": r["id"], "image_url": build_image_url(r["image_path"]), "matiere_id": r["matiere_id"], "material_name": r["nom_matiere"], "reference": r["reference"], "similarity": float(r["similarity"]) if r["similarity"] is not None else None} for r in rows]
            final_results.append({"original_name": original_name, "filename": unique_filename, "url": file_url, "expires_in": "2 hours", "search_results": search_results})
        except Exception as e:
            errors.append(f"{file_ref}: {str(e)}")
    if not final_results and errors:
        return jsonify({"success": False, "message": "All operations failed", "errors": errors}), 500
    return jsonify({"success": True, "message": f"Processed {len(final_results)} files.", "results": final_results, "errors": errors}), 200


# =============================================================================
# MATERIAL DETAILS
# =============================================================================

@app.route("/material_details/<int:matiere_id>", methods=["GET"])
def get_material_details(matiere_id):
    conn = None
    try:
        conn = get_db_conn()
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT * FROM public.matieres WHERE matiere_id = %s", (matiere_id,))
            material = cur.fetchone()
            if not material:
                return jsonify({"success": False, "error": "material_not_found"}), 404
            material = dict(material)
            cur.execute("SELECT fiche_id, date_creation_fiche, derniere_modification FROM public.fiches_matieres WHERE matiere_id = %s ORDER BY fiche_id DESC", (matiere_id,))
            fiches = [dict(row) for row in cur.fetchall()]
            specifications = []
            for fiche in fiches:
                cur.execute("SELECT spec_id, fiche_id, source_type, donnees, date_creation, derniere_modification FROM public.specifications WHERE fiche_id = %s ORDER BY spec_id", (fiche["fiche_id"],))
                specifications.extend([dict(row) for row in cur.fetchall()])
            cur.execute("""
                SELECT men.id, men.matiere_image_id, men.note_json, men.created_at
                FROM public.matiere_expert_notes men
                INNER JOIN public.matiere_images mi ON mi.id = men.matiere_image_id
                WHERE mi.matiere_id = %s ORDER BY men.created_at DESC
            """, (matiere_id,))
            expert_notes = [dict(row) for row in cur.fetchall()]
            return jsonify({"success": True, "material": material, "fiches_matieres": fiches, "specifications": specifications, "expert_notes": expert_notes, "summary": {"matiere_id": matiere_id, "nom_matiere": material.get("nom_matiere"), "reference": material.get("reference"), "type_matiere": material.get("type_matiere"), "num_fiches": len(fiches), "num_specifications": len(specifications), "num_expert_notes": len(expert_notes)}}), 200
    except Exception as e:
        return jsonify({"success": False, "error": "retrieval_failed", "message": str(e)}), 500
    finally:
        if conn:
            conn.close()


# =============================================================================
# FICHE ADN
# =============================================================================

@app.route("/fiche_adn", methods=["GET"])
def get_fiche_adn():
    reference = request.args.get("reference", "").strip()
    if not reference:
        return jsonify({"success": False, "error": "missing_parameters"}), 400
    conn = None
    try:
        conn = get_db_conn()
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT fiche_adn_id, matiere_id, nom_matiere, reference, type_matiere,
                       specifications, num_specifications, date_creation, derniere_modification
                FROM public.fiches_adn_matieres
                WHERE UPPER(REPLACE(TRIM(reference), ' ', '')) = UPPER(REPLACE(%s, ' ', ''))
                LIMIT 1
            """, (reference,))
            result = cur.fetchone()
            if not result:
                return jsonify({"success": False, "error": "fiche_adn_not_found"}), 404
            return jsonify({"success": True, "fiche_adn": dict(result)}), 200
    except Exception as e:
        return jsonify({"success": False, "error": "retrieval_failed", "message": str(e)}), 500
    finally:
        if conn:
            conn.close()


@app.route("/fiche_adn/<int:matiere_id>", methods=["GET"])
def get_fiche_adn_by_id(matiere_id):
    conn = None
    try:
        conn = get_db_conn()
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT fiche_adn_id, matiere_id, nom_matiere, reference, type_matiere,
                       specifications, num_specifications, date_creation, derniere_modification
                FROM public.fiches_adn_matieres WHERE matiere_id = %s LIMIT 1
            """, (matiere_id,))
            result = cur.fetchone()
            if not result:
                return jsonify({"success": False, "error": "fiche_adn_not_found"}), 404
            return jsonify({"success": True, "fiche_adn": dict(result)}), 200
    except Exception as e:
        return jsonify({"success": False, "error": "retrieval_failed", "message": str(e)}), 500
    finally:
        if conn:
            conn.close()


@app.route("/generate_fiche_adn_docx", methods=["GET"])
def generate_fiche_adn_docx():
    reference = request.args.get("reference", "").strip()
    if not reference:
        return jsonify({"success": False, "error": "missing_parameters"}), 400
    conn = None
    try:
        conn = get_db_conn()
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT fiche_adn_id, matiere_id, nom_matiere, reference, type_matiere,
                       specifications, num_specifications, date_creation, derniere_modification
                FROM public.fiches_adn_matieres
                WHERE UPPER(REPLACE(TRIM(reference), ' ', '')) = UPPER(REPLACE(%s, ' ', ''))
                LIMIT 1
            """, (reference,))
            result = cur.fetchone()
            if not result:
                return jsonify({"success": False, "error": "fiche_adn_not_found"}), 404
            result_dict = dict(result)
            matiere_id = result_dict["matiere_id"]
            material_name = result_dict["nom_matiere"]
            type_matiere = result_dict["type_matiere"]
            specifications = result_dict.get("specifications")
            if isinstance(specifications, str):
                try:
                    specifications = json.loads(specifications)
                except:
                    specifications = {}
            if not isinstance(specifications, dict):
                specifications = {}
        content = generate_fiche_adn_content_with_groq(result_dict, material_name, reference, type_matiere, specifications)
        doc = Document()
        title = doc.add_heading(f"MATERIAL DNA SHEET - {material_name}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_paragraph = doc.add_paragraph()
        info_run = info_paragraph.add_run(f"Reference: {reference}\n")
        info_run.bold = True
        info_paragraph.add_run(f"Type: {type_matiere}\n")
        info_paragraph.add_run(f"Generated on: {result_dict.get('date_creation', 'N/A')}")
        doc.add_paragraph()
        doc.add_heading("CONTENT", level=2)
        add_formatted_markdown_to_docx(doc, content)
        images = get_all_images_for_material(matiere_id, limit=10)
        if images:
            doc.add_page_break()
            title_paragraph = doc.add_heading("Examples of micrographie for a reference:", level=2)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            added_magnifications = set()
            images_to_add = []
            for img_data in images:
                if img_data["image_obj"] and len(images_to_add) < 2:
                    magnification = img_data.get("magnification", "N/A")
                    if magnification not in added_magnifications:
                        images_to_add.append(img_data)
                        added_magnifications.add(magnification)
            for idx, img_data in enumerate(images_to_add, 1):
                if img_data["image_obj"]:
                    magnification = img_data.get("magnification", "N/A")
                    mag_heading = doc.add_heading(f"Grossissement: {magnification}x" if magnification != "N/A" else "Grossissement: N/A", level=3)
                    mag_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_stream = io.BytesIO()
                    img_data["image_obj"].save(img_stream, format="PNG")
                    img_stream.seek(0)
                    try:
                        doc.add_picture(img_stream, width=Inches(5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"⚠️ Could not add image {idx}: {e}")
                    doc.add_paragraph()
        timestamp = int(time.time())
        random_id = uuid.uuid4().hex[:8]
        filename = f"Fiche_ADN_{reference}_{timestamp}_{random_id}.docx"
        filepath = DOCX_TEMP_DIR / filename
        doc.save(str(filepath))
        host = request.host or os.getenv("API_HOST", "localhost:5000")
        protocol = request.headers.get("X-Forwarded-Proto", request.scheme)
        if ".azurewebsites.net" in host or ".azure" in host:
            protocol = "https"
        absolute_download_url = f"{protocol}://{host}/download_fiche_adn_docx/{filename}"
        return jsonify({"success": True, "file_name": filename, "download_url": f"/download_fiche_adn_docx/{filename}", "absolute_url": absolute_download_url, "expires_in": "1 hour"}), 200
    except Exception as e:
        return jsonify({"success": False, "error": "generation_failed", "message": str(e)}), 500
    finally:
        if conn:
            conn.close()


@app.route("/download_fiche_adn_docx/<filename>", methods=["GET"])
def download_fiche_adn_docx(filename):
    try:
        if not filename.endswith(".docx") or ".." in filename or "/" in filename or "\\" in filename:
            return jsonify({"success": False, "error": "invalid_file"}), 400
        file_path = DOCX_TEMP_DIR / filename
        if not file_path.exists():
            return jsonify({"success": False, "error": "file_not_found"}), 404
        return send_file(str(file_path), mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name=filename)
    except Exception as e:
        return jsonify({"success": False, "error": "download_failed", "message": str(e)}), 500


# =============================================================================
# DOCX GENERATION — BLACK MIX
# =============================================================================

@app.route("/generate_black_mix_adn_docx", methods=["GET"])
def generate_black_mix_adn_docx():
    mix_id = request.args.get("mix_id", "").strip()
    if not mix_id:
        return jsonify({"success": False, "error": "missing mix_id parameter"}), 400
    conn = None
    try:
        conn = psycopg2.connect(DB_DSN)
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT id, reference, name, status, created_at, document_revision_history FROM public.black_mixes WHERE id = %s", (int(mix_id),))
            bm = cur.fetchone()
            if not bm:
                return jsonify({"success": False, "error": "black_mix_not_found"}), 404
            snapshot = build_black_mix_adn_snapshot(cur, bm["id"], bm["reference"], bm["name"])
            snapshot = serialize_to_json_compatible(snapshot)
        prompt = f"""Tu es un expert en formulation industrielle de matériaux carbone et graphite.
Génère un "Rapport Technique BLACK MIX ADN" COMPLET en suivant cette structure:
#### 1. Identité du Black Mix
#### 2. Composition et ADN des Composants
#### 3. Processus de Fabrication (Mischkarte)
#### 4. Plan de Contrôle
#### 5. Synthèse de l'Identité Structurelle
RÈGLES: Aucune hallucination. Langue: Français. Style professionnel.
### DONNÉES SOURCE (JSON):
{json.dumps(snapshot, indent=2, ensure_ascii=False, default=str)}"""
        ai_content = call_groq_with_retry(
            messages=[{"role": "system", "content": "Tu es un expert en formulation industrielle."},
                      {"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile", temperature=0.2, max_tokens=6000
        )
        content = ai_content.choices[0].message.content if ai_content.choices else ""
        doc = Document()
        title = doc.add_heading(f"RAPPORT ADN — BLACK MIX {bm['name']}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info = doc.add_paragraph()
        info.add_run(f"Référence: {bm['reference']}\n").bold = True
        info.add_run(f"Statut: {bm['status']}\n")
        info.add_run(f"Généré le: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph()
        add_formatted_markdown_to_docx(doc, content)
        timestamp = int(time.time())
        random_id = uuid.uuid4().hex[:8]
        filename = f"ADN_BlackMix_{bm['reference']}_{timestamp}_{random_id}.docx"
        filepath = DOCX_TEMP_DIR / filename
        doc.save(str(filepath))
        host = request.host or os.getenv("API_HOST", "localhost:5000")
        protocol = request.headers.get("X-Forwarded-Proto", request.scheme)
        if ".azurewebsites.net" in host or ".azure" in host:
            protocol = "https"
        absolute_url = f"{protocol}://{host}/download_fiche_adn_docx/{filename}"
        return jsonify({"success": True, "file_name": filename, "download_url": f"/download_fiche_adn_docx/{filename}", "absolute_url": absolute_url, "expires_in": "1 hour"}), 200
    except Exception as e:
        return jsonify({"success": False, "error": "generation_failed", "message": str(e)}), 500
    finally:
        if conn:
            conn.close()


# =============================================================================
# DOCX GENERATION — NUANCE
# =============================================================================

@app.route("/generate_nuance_adn_docx", methods=["GET"])
def generate_nuance_adn_docx():
    nuance_id = request.args.get("nuance_id", "").strip()
    if not nuance_id:
        return jsonify({"success": False, "error": "missing nuance_id parameter"}), 400
    conn = None
    try:
        conn = psycopg2.connect(DB_DSN)
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT id, reference, name, status, created_at, document_revision_history FROM public.nuances WHERE id = %s", (int(nuance_id),))
            nuance = cur.fetchone()
            if not nuance:
                return jsonify({"success": False, "error": "nuance_not_found"}), 404
            snapshot = build_nuance_adn_snapshot(cur, nuance["id"], nuance["reference"], nuance["name"])
            snapshot = serialize_to_json_compatible(snapshot)
            for img in snapshot.get("images", []):
                if img.get("image_path"):
                    try:
                        img["image_url"] = build_image_url(img["image_path"])
                    except RuntimeError:
                        img["image_url"] = None
        prompt = f"""Tu es un expert en formulation industrielle de matériaux carbone et graphite.
Génère un "Rapport Technique NUANCE ADN" COMPLET en suivant cette structure:
#### 1. Identité de la Nuance
#### 2. Programme de Cuisson (Wärme-Nachbehandlung)
#### 3. Composition et ADN des Composants
#### 4. Processus de Fabrication (Mischkarte)
#### 5. Plan de Contrôle
#### 6. Historique des Révisions
#### 7. Synthèse de l'Identité Structurelle
RÈGLES: Aucune hallucination. Langue: Français. Style professionnel.
### DONNÉES SOURCE (JSON):
{json.dumps(snapshot, indent=2, ensure_ascii=False, default=str)}"""
        ai_content = call_groq_with_retry(
            messages=[{"role": "system", "content": "Tu es un expert en formulation industrielle."},
                      {"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile", temperature=0.2, max_tokens=6000
        )
        content = ai_content.choices[0].message.content if ai_content.choices else ""
        doc = Document()
        title = doc.add_heading(f"RAPPORT ADN — NUANCE {nuance['name']}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info = doc.add_paragraph()
        info.add_run(f"Référence: {nuance['reference']}\n").bold = True
        info.add_run(f"Statut: {nuance['status']}\n")
        info.add_run(f"Généré le: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph()
        doc.add_heading("CONTENU", level=2)
        add_formatted_markdown_to_docx(doc, content)
        images = snapshot.get("images", [])
        if images:
            doc.add_page_break()
            img_title = doc.add_heading("Micrographies de la Nuance", level=2)
            img_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for img_data in images[:4]:
                image_path = img_data.get("image_path")
                if not image_path:
                    continue
                normalized = image_path.replace("\\\\", "/")
                fname = Path(normalized).name
                img_obj = None
                for fpath in [Path(normalized), BASE_DIR / normalized, IMAGES_DIR / fname, BASE_DIR / "output_v3" / "images" / fname, BASE_DIR / "output_v4" / "images" / fname]:
                    try:
                        if fpath.exists():
                            img_obj = Image.open(fpath).convert("RGB")
                            break
                    except Exception:
                        pass
                if img_obj:
                    img_stream = io.BytesIO()
                    img_obj.save(img_stream, format="PNG")
                    img_stream.seek(0)
                    try:
                        doc.add_picture(img_stream, width=Inches(5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"⚠️ Could not add image: {e}")
                    doc.add_paragraph()
        timestamp = int(time.time())
        random_id = uuid.uuid4().hex[:8]
        filename = f"ADN_Nuance_{nuance['reference']}_{timestamp}_{random_id}.docx"
        filepath = DOCX_TEMP_DIR / filename
        doc.save(str(filepath))
        host = request.host or os.getenv("API_HOST", "localhost:5000")
        protocol = request.headers.get("X-Forwarded-Proto", request.scheme)
        if ".azurewebsites.net" in host or ".azure" in host:
            protocol = "https"
        absolute_url = f"{protocol}://{host}/download_fiche_adn_docx/{filename}"
        return jsonify({"success": True, "file_name": filename, "download_url": f"/download_fiche_adn_docx/{filename}", "absolute_url": absolute_url, "expires_in": "1 hour"}), 200
    except Exception as e:
        return jsonify({"success": False, "error": "generation_failed", "message": str(e)}), 500
    finally:
        if conn:
            conn.close()


@app.route("/populate_fiches_adn_table", methods=["POST"])
def populate_fiches_adn_table():
    try:
        conn = get_db_conn()
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT matiere_id, nom_matiere, reference, type_matiere FROM public.matieres ORDER BY matiere_id")
            materials = cur.fetchall()
            if not materials:
                return jsonify({"success": False, "message": "No materials found"}), 404
            processed_count = updated_count = inserted_count = error_count = 0
            errors_details = []
            for material in materials:
                try:
                    matiere_id = material["matiere_id"]
                    cur.execute("SELECT fiche_id FROM public.fiches_matieres WHERE matiere_id = %s ORDER BY fiche_id DESC", (matiere_id,))
                    fiches = [dict(row) for row in cur.fetchall()]
                    specifications_list = []
                    for fiche in fiches:
                        cur.execute("SELECT spec_id, fiche_id, source_type, donnees, date_creation, derniere_modification FROM public.specifications WHERE fiche_id = %s ORDER BY spec_id", (fiche["fiche_id"],))
                        specifications_list.extend([dict(row) for row in cur.fetchall()])
                    cur.execute("""
                        SELECT men.id, men.matiere_image_id, men.note_json, men.created_at, mi.image_path
                        FROM public.matiere_expert_notes men
                        INNER JOIN public.matiere_images mi ON mi.id = men.matiere_image_id
                        WHERE mi.matiere_id = %s ORDER BY men.created_at DESC
                    """, (matiere_id,))
                    expert_notes = [dict(row) for row in cur.fetchall()]
                    aggregated_data = serialize_to_json_compatible({"fiches": fiches, "specifications": specifications_list, "expert_notes": expert_notes, "summary": {"num_fiches": len(fiches), "num_specifications": len(specifications_list), "num_expert_notes": len(expert_notes)}})
                    cur.execute("SELECT fiche_adn_id FROM public.fiches_adn_matieres WHERE matiere_id = %s", (matiere_id,))
                    existing = cur.fetchone()
                    if existing:
                        cur.execute("UPDATE public.fiches_adn_matieres SET nom_matiere=%s, reference=%s, type_matiere=%s, specifications=%s, num_specifications=%s, derniere_modification=CURRENT_TIMESTAMP WHERE matiere_id=%s",
                                    (material["nom_matiere"], material["reference"], material["type_matiere"], Json(aggregated_data), len(specifications_list), matiere_id))
                        updated_count += 1
                    else:
                        cur.execute("INSERT INTO public.fiches_adn_matieres (matiere_id, nom_matiere, reference, type_matiere, specifications, num_specifications, date_creation, derniere_modification) VALUES (%s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)",
                                    (matiere_id, material["nom_matiere"], material["reference"], material["type_matiere"], Json(aggregated_data), len(specifications_list)))
                        inserted_count += 1
                    processed_count += 1
                    conn.commit()
                except Exception as mat_error:
                    error_count += 1
                    errors_details.append({"matiere_id": material.get("matiere_id"), "error": str(mat_error)})
                    conn.rollback()
            return jsonify({"success": error_count == 0, "summary": {"total_materials": len(materials), "processed": processed_count, "inserted": inserted_count, "updated": updated_count, "errors": error_count}, "errors_details": errors_details[:10]}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        if conn:
            conn.close()


@app.route("/verify_fiches_adn_table", methods=["GET"])
def verify_fiches_adn_table():
    try:
        conn = get_db_conn()
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT COUNT(*) as total FROM public.fiches_adn_matieres")
            total_count = cur.fetchone()["total"]
            if total_count == 0:
                return jsonify({"success": True, "message": "Table is empty", "total_records": 0}), 200
            cur.execute("""
                SELECT COUNT(*) as total_records, SUM(num_specifications) as total_specifications,
                       AVG(num_specifications) as avg_specifications_per_material,
                       MIN(num_specifications) as min_specifications, MAX(num_specifications) as max_specifications,
                       COUNT(CASE WHEN num_specifications = 0 THEN 1 END) as materials_without_specs,
                       COUNT(CASE WHEN specifications IS NOT NULL THEN 1 END) as materials_with_data
                FROM public.fiches_adn_matieres
            """)
            stats = dict(cur.fetchone())
            cur.execute("SELECT fiche_adn_id, matiere_id, nom_matiere, reference, type_matiere, num_specifications, date_creation, derniere_modification, specifications FROM public.fiches_adn_matieres ORDER BY derniere_modification DESC LIMIT 5")
            samples = []
            for row in cur.fetchall():
                sample = dict(row)
                specs_data = sample.get("specifications", {})
                if isinstance(specs_data, str):
                    try:
                        specs_data = json.loads(specs_data)
                    except:
                        specs_data = {}
                sample["specifications_summary"] = {"num_fiches": specs_data.get("summary", {}).get("num_fiches", 0) if isinstance(specs_data, dict) else 0, "num_specifications": specs_data.get("summary", {}).get("num_specifications", 0) if isinstance(specs_data, dict) else 0, "num_expert_notes": specs_data.get("summary", {}).get("num_expert_notes", 0) if isinstance(specs_data, dict) else 0}
                del sample["specifications"]
                samples.append(sample)
            return jsonify({"success": True, "total_records": total_count, "statistics": stats, "samples": samples}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        if conn:
            conn.close()


@app.route("/search", methods=["POST"])
def search():
    data = request.get_json(silent=True) or {}
    if not data:
        return jsonify({"success": False, "error": "missing_json_body"}), 400
    top_k = int(data.get("top_k", 5))
    if top_k < 1 or top_k > 50:
        return jsonify({"success": False, "error": "invalid_top_k"}), 400
    temp_filename = data.get("temp_filename")
    file_id = data.get("file_id")
    download_link = data.get("download_link")
    provided = [bool(download_link), bool(temp_filename), bool(file_id)]
    if sum(provided) != 1:
        return jsonify({"success": False, "error": "Provide exactly ONE of: download_link, temp_filename, file_id"}), 400
    img = None
    if download_link:
        try:
            r = requests.get(download_link, timeout=20)
            r.raise_for_status()
            img = Image.open(io.BytesIO(r.content)).convert("RGB")
        except Exception as e:
            return jsonify({"success": False, "error": "download_link_failed", "message": str(e)}), 400
    elif temp_filename:
        file_path = TEMP_UPLOAD_DIR / temp_filename
        if not file_path.exists():
            return jsonify({"success": False, "error": "temp_file_not_found"}), 404
        img = Image.open(file_path).convert("RGB")
    elif file_id:
        if not client:
            return jsonify({"success": False, "error": "openai_not_configured"}), 400
        file_content = client.files.content(file_id).read()
        img = Image.open(io.BytesIO(file_content)).convert("RGB")
    try:
        query_embedding = compute_embedding_from_pil(img)
        rows = search_similar_in_db(query_embedding, top_k=top_k)
        results = [{"id": r["id"], "image_url": build_image_url(r["image_path"]), "matiere_id": r["matiere_id"], "material_name": r["nom_matiere"], "reference": r["reference"], "similarity": float(r["similarity"]) if r["similarity"] is not None else None} for r in rows]
        return jsonify({"success": True, "results": results}), 200
    except Exception as e:
        return jsonify({"success": False, "error": "search_failed", "message": str(e)}), 500


# =============================================================================
# APPLICATION ANALYSIS
# =============================================================================

def generate_application_analysis_with_llm(fiche_adn_data, company_context=None):
    if not groq_client:
        return {"success": False, "error": "llm_not_configured"}
    if not company_context:
        company_context = """AVOCarbon - Company Scope:
Core Business Areas:
1. Carbon brushes for electric motors
2. Brush-holder assemblies
3. Inductors and coils (chokes) for EMI filtering
4. Dynamic sealing joints (via Cyclam division)
5. Self-lubricating bearings and bushings
6. Friction rings, rotors, vanes for motors and pumps
Target Industries: Automotive, Power tools, Household appliances, Industrial equipment"""
    material_name = fiche_adn_data.get("nom_matiere", "Unknown Material")
    reference = fiche_adn_data.get("reference", "N/A")
    type_matiere = fiche_adn_data.get("type_matiere", "N/A")
    specifications = fiche_adn_data.get("specifications", {})
    prompt = f"""# Material Application Analysis Request
Material: {material_name} | Reference: {reference} | Type: {type_matiere}
Company Context: {company_context}
Specifications: {json.dumps(specifications, indent=2, ensure_ascii=False)}
Provide a comprehensive JSON analysis with this structure:
{{"material_summary": {{"key_characteristics": [], "primary_domains": []}}, "applications": [{{"application_name": "string", "application_category": "string", "industry_sector": "string", "domain": "core_business | strategic_opportunity | outside_scope", "priority_level": 1, "engagement_process": {{"process_description": "string", "steps": [], "material_role": "string"}}, "required_properties": [{{"property_name": "string", "importance": "critical", "reason": "string"}}], "market_potential": {{"growth_trend": "growing", "competitive_advantage": "string"}}}}], "strategic_recommendations": {{"within_scope": [{{"opportunity": "string", "rationale": "string", "development_effort": "low"}}], "strategic_expansion": [{{"opportunity": "string", "rationale": "string", "requirements": "string"}}]}}}}
Respond ONLY with valid JSON."""
    try:
        response = call_groq_with_retry(messages=[{"role": "system", "content": "You are an expert materials engineer. Respond only with valid JSON."}, {"role": "user", "content": prompt}], model="llama-3.3-70b-versatile", temperature=0.3, max_tokens=8000, response_format={"type": "json_object"})
        analysis_data = json.loads(response.choices[0].message.content)
        return {"success": True, "analysis": analysis_data, "model_used": "llama-3.3-70b-versatile"}
    except json.JSONDecodeError as e:
        return {"success": False, "error": "json_parse_error", "message": str(e)}
    except Exception as e:
        return {"success": False, "error": "llm_generation_failed", "message": str(e)}


def generate_application_analysis_docx_with_llm(fiche_data, analysis_data):
    if not groq_client:
        raise Exception("Groq client not initialized")
    prompt = f"""Tu es un expert en formulation industrielle. Génère le contenu markdown complet d'un rapport d'analyse d'usage pour:
- Matière: {fiche_data.get('nom_matiere', 'N/A')} (Ref: {fiche_data.get('reference', 'N/A')})
Données d'analyse: {json.dumps(analysis_data, indent=2, ensure_ascii=False)}
Structure requise:
### 1) Lecture rapide du matériau
### 2) Domaines d'application principaux (A, B, C...)
### 3) Tableau de synthèse
### 4) Applications stratégiques hors cœur de métier
### 5) Lecture stratégique pour votre groupe
Génère uniquement le contenu markdown."""
    response = call_groq_with_retry(messages=[{"role": "system", "content": "You are a technical writer generating markdown content for reports."}, {"role": "user", "content": prompt}], model="llama-3.3-70b-versatile", temperature=0.2, max_tokens=4000)
    markdown_content = response.choices[0].message.content
    doc = Document()
    doc.add_heading(f"Analyse d'usage du {fiche_data.get('nom_matiere', 'Matière')} {fiche_data.get('reference', 'N/A')}", level=1)
    add_formatted_markdown_to_docx(doc, markdown_content)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    ref_safe = fiche_data.get("reference", "material").replace(" ", "_")
    filename = f"Analyse_{ref_safe}_{timestamp}.docx"
    filepath = DOCX_TEMP_DIR / filename
    doc.save(str(filepath))
    return filename


@app.route("/generate_application_analysis", methods=["POST"])
def generate_application_analysis():
    data = request.get_json() or {}
    reference = data.get("reference", "").strip()
    company_context = data.get("company_context")
    save_to_db = data.get("save_to_db", True)
    if not reference:
        return jsonify({"success": False, "error": "missing_parameters"}), 400
    conn = None
    try:
        conn = get_db_conn()
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT fiche_adn_id, matiere_id, nom_matiere, reference, type_matiere, specifications FROM public.fiches_adn_matieres WHERE UPPER(REPLACE(TRIM(reference), ' ', '')) = UPPER(REPLACE(%s, ' ', '')) LIMIT 1", (reference,))
            fiche_adn = cur.fetchone()
            if not fiche_adn:
                return jsonify({"success": False, "error": "fiche_adn_not_found"}), 404
            fiche_data = dict(fiche_adn)
            cur.execute("SELECT fiche_app_id, analysis_data FROM public.fiches_applications_matieres WHERE UPPER(REPLACE(TRIM(reference), ' ', '')) = UPPER(REPLACE(%s, ' ', '')) LIMIT 1", (reference,))
            existing = cur.fetchone()
            if existing:
                existing_dict = dict(existing)
                existing_analysis = existing_dict.get("analysis_data", {})
                docx_filename = generate_application_analysis_docx_with_llm(fiche_data, existing_analysis)
                protocol = request.headers.get("X-Forwarded-Proto", request.scheme)
                if ".azurewebsites.net" in request.host or ".azure" in request.host:
                    protocol = "https"
                download_url = url_for('download_fiche_adn_docx', filename=docx_filename, _external=True, _scheme=protocol)
                return jsonify({"success": True, "message": "Analysis already exists", "analysis": existing_analysis, "download_url": download_url, "is_existing": True}), 200
        analysis_result = generate_application_analysis_with_llm(fiche_data, company_context)
        if not analysis_result.get("success"):
            return jsonify(analysis_result), 500
        if save_to_db:
            analysis_data = analysis_result["analysis"]
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("SELECT fiche_adn_id FROM public.fiches_adn_matieres WHERE matiere_id = %s LIMIT 1", (fiche_data["matiere_id"],))
                fiche_adn_row = cur.fetchone()
                fiche_adn_id = fiche_adn_row["fiche_adn_id"] if fiche_adn_row else None
                cur.execute("INSERT INTO public.fiches_applications_matieres (matiere_id, fiche_adn_id, nom_matiere, reference, type_matiere, analysis_data, num_applications, date_creation, derniere_modification) VALUES (%s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP) RETURNING fiche_app_id",
                            (fiche_data["matiere_id"], fiche_adn_id, fiche_data["nom_matiere"], fiche_data["reference"], fiche_data["type_matiere"], Json(analysis_data), len(analysis_data.get("applications", []))))
                fiche_app_id = cur.fetchone()["fiche_app_id"]
                conn.commit()
                analysis_result["fiche_app_id"] = fiche_app_id
                analysis_result["saved_to_database"] = True
        docx_filename = generate_application_analysis_docx_with_llm(fiche_data, analysis_result["analysis"])
        protocol = request.headers.get("X-Forwarded-Proto", request.scheme)
        if ".azurewebsites.net" in request.host or ".azure" in request.host:
            protocol = "https"
        download_url = url_for('download_fiche_adn_docx', filename=docx_filename, _external=True, _scheme=protocol)
        analysis_result["docx_filename"] = docx_filename
        analysis_result["download_url"] = download_url
        return jsonify(analysis_result), 200
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"success": False, "error": "analysis_generation_failed", "message": str(e)}), 500
    finally:
        if conn:
            conn.close()


@app.route("/application_analysis", methods=["GET"])
def get_application_analysis():
    reference = request.args.get("reference", "").strip()
    include_sessions = request.args.get("include_sessions", "false").lower() == "true"
    include_steps = request.args.get("include_steps", "true").lower() == "true"
    if not reference:
        return jsonify({"success": False, "error": "missing_parameters"}), 400
    conn = None
    try:
        conn = get_db_conn()
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT m.matiere_id, m.nom_matiere, m.reference, m.type_matiere FROM public.matieres m WHERE UPPER(REPLACE(TRIM(m.reference), ' ', '')) = UPPER(REPLACE(%s, ' ', '')) LIMIT 1", (reference,))
            material = cur.fetchone()
            if not material:
                return jsonify({"success": False, "error": "material_not_found"}), 404
            material = dict(material)
            matiere_id = material["matiere_id"]
            cur.execute("SELECT fiche_app_id, fiche_adn_id, analysis_data, num_applications, date_creation, derniere_modification FROM public.fiches_applications_matieres WHERE matiere_id = %s ORDER BY date_creation DESC", (matiere_id,))
            fiches = [dict(row) for row in cur.fetchall()]
            latest_analysis = fiches[0] if fiches else None
            applications = []
            if latest_analysis:
                analysis_data = latest_analysis.get("analysis_data", {})
                applications = analysis_data.get("applications", [])
            if include_steps and applications:
                for app in applications:
                    process = app.get("engagement_process", {})
                    app["process_steps"] = process.get("steps", [])
            summary = {"total_applications": len(applications), "by_domain": {}, "by_priority": {"high": 0, "medium": 0, "low": 0}, "total_analyses": len(fiches)}
            for app in applications:
                domain = app.get("domain", "unknown")
                summary["by_domain"][domain] = summary["by_domain"].get(domain, 0) + 1
                priority = app.get("priority_level", 0)
                if priority >= 4:
                    summary["by_priority"]["high"] += 1
                elif priority >= 2:
                    summary["by_priority"]["medium"] += 1
                else:
                    summary["by_priority"]["low"] += 1
            docx_url = None
            if latest_analysis:
                fiche_data = dict(material)
                docx_filename = generate_application_analysis_docx_with_llm(fiche_data, latest_analysis.get("analysis_data", {}))
                protocol = request.headers.get("X-Forwarded-Proto", request.scheme)
                if ".azurewebsites.net" in request.host or ".azure" in request.host:
                    protocol = "https"
                docx_url = url_for('download_fiche_adn_docx', filename=docx_filename, _external=True, _scheme=protocol)
            return jsonify({"success": True, "material": material, "applications": applications, "analysis_sessions": fiches if include_sessions else None, "summary": summary, "docx_download_url": docx_url}), 200
    except Exception as e:
        return jsonify({"success": False, "error": "retrieval_failed", "message": str(e)}), 500
    finally:
        if conn:
            conn.close()


# =============================================================================
# BLACK MIX — CORE HELPERS
# =============================================================================

def resolve_ref_lookup(cur, components):
    ref_lookup = {}
    validation_errors = []
    for component in components:
        ref = component.get("reference")
        if not ref:
            validation_errors.append("Component missing reference")
            continue
        if ref in ref_lookup:
            continue
        cur.execute("SELECT id FROM public.black_mixes WHERE reference = %s LIMIT 1", (ref,))
        bm_row = cur.fetchone()
        if bm_row:
            ref_lookup[ref] = {"type": "black_mix", "id": bm_row[0]}
            continue
        cur.execute("SELECT matiere_id FROM public.matieres WHERE reference = %s LIMIT 1", (ref,))
        mat_row = cur.fetchone()
        if mat_row:
            ref_lookup[ref] = {"type": "matiere", "id": mat_row[0]}
            continue
        validation_errors.append(f"Reference '{ref}' not found in black_mixes or matieres")
    return ref_lookup, validation_errors


def build_black_mix_adn_snapshot(cur, black_mix_id, product_reference, mix_name, _visited=None):
    if _visited is None:
        _visited = set()
    if black_mix_id in _visited:
        return {"black_mix_id": black_mix_id, "product_reference": product_reference, "mix_name": mix_name, "error": "Circular reference detected — snapshot truncated here"}
    _visited.add(black_mix_id)
    cur.execute("SELECT document_revision_history FROM public.black_mixes WHERE id = %s", (black_mix_id,))
    bm_row = cur.fetchone()
    revision_history = bm_row["document_revision_history"] if bm_row and bm_row["document_revision_history"] else None
    cur.execute("""
        SELECT c.id, c.component_name, c.quantity_value, c.quantity_unit, c.metadata,
               c.matiere_id, c.sub_black_mix_id,
               m.reference AS matiere_reference, m.nom_matiere AS matiere_name,
               bm.reference AS sub_bm_reference, bm.name AS sub_bm_name
        FROM public.black_mix_components c
        LEFT JOIN public.matieres m ON m.matiere_id = c.matiere_id
        LEFT JOIN public.black_mixes bm ON bm.id = c.sub_black_mix_id
        WHERE c.black_mix_id = %s ORDER BY c.id
    """, (black_mix_id,))
    raw_components = cur.fetchall()
    components = []
    for r in raw_components:
        is_sub = r["sub_black_mix_id"] is not None
        entry = {"id": r["id"], "component_name": r["component_name"], "quantity": float(r["quantity_value"]) if r["quantity_value"] is not None else None, "unit": r["quantity_unit"], "metadata": r["metadata"], "is_sub_black_mix": is_sub, "reference": r["sub_bm_reference"] if is_sub else r["matiere_reference"], "material_name": r["sub_bm_name"] if is_sub else r["matiere_name"], "sub_black_mix_id": r["sub_black_mix_id"], "matiere_id": r["matiere_id"], "sub_black_mix_adn": None}
        if is_sub:
            entry["sub_black_mix_adn"] = build_black_mix_adn_snapshot(cur, r["sub_black_mix_id"], r["sub_bm_reference"], r["sub_bm_name"], _visited=set(_visited))
        components.append(entry)
    cur.execute("SELECT s.id, s.step_order, s.step_name, s.machine_name, s.parameters FROM public.black_mix_process_steps s WHERE s.black_mix_id = %s ORDER BY s.step_order", (black_mix_id,))
    steps_raw = cur.fetchall()
    process_steps = []
    for s in steps_raw:
        cur.execute("""
            SELECT sm.matiere_id, sm.sub_black_mix_id, m.reference AS matiere_ref, bm.reference AS sub_bm_ref
            FROM public.black_mix_step_materials sm
            LEFT JOIN public.matieres m ON m.matiere_id = sm.matiere_id
            LEFT JOIN public.black_mixes bm ON bm.id = sm.sub_black_mix_id
            WHERE sm.process_step_id = %s
        """, (s["id"],))
        step_mat_refs = [sm["sub_bm_ref"] if sm["sub_black_mix_id"] is not None else sm["matiere_ref"] for sm in cur.fetchall() if (sm["sub_bm_ref"] if sm["sub_black_mix_id"] is not None else sm["matiere_ref"])]
        process_steps.append({"step_order": s["step_order"], "step_name": s["step_name"], "machine": s["machine_name"], "parameters": s["parameters"], "materials": step_mat_refs})
    cur.execute("SELECT parameter_name, target_value, min_value, max_value, unit FROM public.black_mix_control_plan WHERE black_mix_id = %s ORDER BY parameter_name", (black_mix_id,))
    control_plan = [{"parameter_name": r["parameter_name"], "target_value": float(r["target_value"]) if r["target_value"] is not None else None, "min_value": float(r["min_value"]) if r["min_value"] is not None else None, "max_value": float(r["max_value"]) if r["max_value"] is not None else None, "unit": r["unit"]} for r in cur.fetchall()]

    def flatten(comp_list, depth=0):
        flat = []
        for c in comp_list:
            flat.append({**c, "depth": depth, "sub_black_mix_adn": None})
            if c["is_sub_black_mix"] and c["sub_black_mix_adn"]:
                flat.extend(flatten(c["sub_black_mix_adn"].get("composition", []), depth + 1))
        return flat

    return {"black_mix_id": black_mix_id, "product_reference": product_reference, "mix_name": mix_name, "status": "draft", "document_revision_history": revision_history, "created_at": datetime.now().isoformat(), "composition": components, "composition_flat": flatten(components), "process_steps": process_steps, "step_materials": {str(s["step_order"]): s["materials"] for s in process_steps}, "control_plan": control_plan, "snapshot_timestamp": datetime.now().isoformat()}


# =============================================================================
# BLACK MIX — ENDPOINTS
# =============================================================================

@app.route("/black-mix/validate-material/<reference>", methods=["GET"])
def validate_black_mix_material(reference):
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, name FROM public.black_mixes WHERE reference = %s LIMIT 1", (reference,))
            bm_row = cur.fetchone()
            if bm_row:
                return jsonify({"reference": reference, "exists": True, "component_type": "black_mix", "material_name": bm_row[1], "id": bm_row[0]}), 200
            cur.execute("SELECT matiere_id, nom_matiere FROM public.matieres WHERE reference = %s LIMIT 1", (reference,))
            mat_row = cur.fetchone()
            if mat_row:
                return jsonify({"reference": reference, "exists": True, "component_type": "matiere", "material_name": mat_row[1], "id": mat_row[0]}), 200
            return jsonify({"reference": reference, "exists": False, "component_type": None, "material_name": None, "id": None}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/black-mix/submit", methods=["POST"])
def submit_black_mix():
    if not request.is_json:
        return jsonify({"success": False, "error": "Request body must be JSON"}), 400
    data = request.get_json()
    product_reference = data.get("product_reference")
    mix_name = data.get("mix_name")
    components = data.get("components", [])
    process_steps = data.get("process_steps", [])
    step_materials_map = data.get("step_materials", {})
    control_plan = data.get("control_plan", [])
    document_revision_history = data.get("document_revision_history")
    if not product_reference or not mix_name:
        return jsonify({"success": False, "error": "product_reference and mix_name are required"}), 400
    if not process_steps:
        return jsonify({"success": False, "error": "At least one process_step is required"}), 400
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn:
            with conn.cursor() as cur:
                ref_lookup, validation_errors = resolve_ref_lookup(cur, components)
                if validation_errors:
                    return jsonify({"success": False, "validation_errors": validation_errors}), 400
                all_step_refs = {ref for refs in step_materials_map.values() for ref in refs if ref}
                extra_refs = all_step_refs - set(ref_lookup.keys())
                if extra_refs:
                    extra_lookup, extra_errors = resolve_ref_lookup(cur, [{"reference": r} for r in extra_refs])
                    if extra_errors:
                        return jsonify({"success": False, "validation_errors": extra_errors}), 400
                    ref_lookup.update(extra_lookup)
                cur.execute("INSERT INTO public.black_mixes (reference, name, status, created_at, document_revision_history) VALUES (%s, %s, 'draft', NOW(), %s) RETURNING id",
                            (product_reference, mix_name, Json(document_revision_history) if document_revision_history else None))
                black_mix_id = cur.fetchone()[0]
                for component in components:
                    ref = component.get("reference")
                    resolved = ref_lookup[ref]
                    cur.execute("INSERT INTO public.black_mix_components (black_mix_id, matiere_id, sub_black_mix_id, component_name, quantity_value, quantity_unit, metadata) VALUES (%s, %s, %s, %s, %s, %s, %s)",
                                (black_mix_id, resolved["id"] if resolved["type"] == "matiere" else None, resolved["id"] if resolved["type"] == "black_mix" else None, component.get("component_name") or ref, component.get("quantity"), component.get("unit", "kg"), Json(component.get("metadata", {}))))
                for step in process_steps:
                    step_order = step.get("step_order")
                    cur.execute("INSERT INTO public.black_mix_process_steps (black_mix_id, step_order, step_name, machine_name, parameters) VALUES (%s, %s, %s, %s, %s) RETURNING id",
                                (black_mix_id, step_order, step.get("step_name"), step.get("machine"), Json(step.get("parameters", {}))))
                    process_step_id = cur.fetchone()[0]
                    refs_for_step = step_materials_map.get(str(step_order), [])
                    if not refs_for_step:
                        raise ValueError(f"Step '{step.get('step_name')}' (order {step_order}) has no materials in step_materials")
                    seen_ids = set()
                    for ref in refs_for_step:
                        if not ref:
                            continue
                        resolved = ref_lookup.get(ref)
                        if not resolved:
                            raise ValueError(f"Reference '{ref}' in step_materials not found")
                        dedup_key = (resolved["type"], resolved["id"])
                        if dedup_key in seen_ids:
                            continue
                        seen_ids.add(dedup_key)
                        if resolved["type"] == "matiere":
                            cur.execute("INSERT INTO public.black_mix_step_materials (process_step_id, matiere_id, created_at) VALUES (%s, %s, NOW())", (process_step_id, resolved["id"]))
                        else:
                            cur.execute("INSERT INTO public.black_mix_step_materials (process_step_id, sub_black_mix_id, created_at) VALUES (%s, %s, NOW())", (process_step_id, resolved["id"]))
                for param in control_plan:
                    cur.execute("INSERT INTO public.black_mix_control_plan (black_mix_id, parameter_name, target_value, min_value, max_value, unit) VALUES (%s, %s, %s, %s, %s, %s)",
                                (black_mix_id, param.get("parameter_name"), param.get("target_value"), param.get("min_value"), param.get("max_value"), param.get("unit")))
                adn_snapshot = build_black_mix_adn_snapshot(cur, black_mix_id, product_reference, mix_name)
                cur.execute("INSERT INTO public.black_mix_adn (black_mix_id, adn_text, version, created_at) VALUES (%s, %s, 1, NOW()) RETURNING id", (black_mix_id, Json(adn_snapshot)))
                adn_id = cur.fetchone()[0]
                return jsonify({"success": True, "message": f"Black Mix '{mix_name}' created successfully", "black_mix_id": black_mix_id, "product_reference": product_reference, "component_types": {ref: info["type"] for ref, info in ref_lookup.items()}, "adn": {"id": adn_id, "version": 1}}), 200
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/black-mix/list", methods=["GET"])
def list_black_mixes():
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, reference, name, status, created_at FROM public.black_mixes ORDER BY created_at DESC")
            rows = cur.fetchall()
            return jsonify({"success": True, "black_mixes": [{"id": r[0], "product_reference": r[1], "mix_name": r[2], "status": r[3], "created_at": r[4].isoformat() if r[4] else None} for r in rows]}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/black-mix/<int:mix_id>", methods=["GET"])
def get_black_mix_details(mix_id):
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, reference, name, status, created_at, document_revision_history FROM public.black_mixes WHERE id = %s", (mix_id,))
            row = cur.fetchone()
            if not row:
                return jsonify({"success": False, "error": "Black Mix not found"}), 404
            result = {"id": row[0], "product_reference": row[1], "mix_name": row[2], "status": row[3], "created_at": row[4].isoformat() if row[4] else None, "document_revision_history": row[5]}
            cur.execute("""
                SELECT c.id, c.component_name, c.quantity_value, c.quantity_unit, c.matiere_id, c.sub_black_mix_id, c.metadata,
                       m.reference AS matiere_ref, m.nom_matiere AS matiere_name,
                       bm.reference AS sub_bm_ref, bm.name AS sub_bm_name
                FROM public.black_mix_components c
                LEFT JOIN public.matieres m ON m.matiere_id = c.matiere_id
                LEFT JOIN public.black_mixes bm ON bm.id = c.sub_black_mix_id
                WHERE c.black_mix_id = %s ORDER BY c.id
            """, (mix_id,))
            result["components"] = []
            for r in cur.fetchall():
                is_sub = r[5] is not None
                result["components"].append({"id": r[0], "component_name": r[1], "quantity": float(r[2]) if r[2] is not None else None, "unit": r[3], "is_sub_black_mix": is_sub, "reference": r[9] if is_sub else r[7], "material_name": r[10] if is_sub else r[8], "matiere_id": r[4], "sub_black_mix_id": r[5], "metadata": r[6]})
            cur.execute("""
                SELECT s.id, s.step_order, s.step_name, s.machine_name, s.parameters,
                       ARRAY_AGG(COALESCE(bm.reference, m.reference) ORDER BY COALESCE(bm.reference, m.reference)) AS materials
                FROM public.black_mix_process_steps s
                LEFT JOIN public.black_mix_step_materials sm ON sm.process_step_id = s.id
                LEFT JOIN public.matieres m ON m.matiere_id = sm.matiere_id
                LEFT JOIN public.black_mixes bm ON bm.id = sm.sub_black_mix_id
                WHERE s.black_mix_id = %s GROUP BY s.id ORDER BY s.step_order
            """, (mix_id,))
            result["process_steps"] = [{"step_order": r[1], "step_name": r[2], "machine": r[3], "parameters": r[4], "materials": [x for x in (r[5] or []) if x is not None]} for r in cur.fetchall()]
            cur.execute("SELECT parameter_name, target_value, min_value, max_value, unit FROM public.black_mix_control_plan WHERE black_mix_id = %s", (mix_id,))
            result["control_plan"] = [{"parameter_name": r[0], "target_value": float(r[1]) if r[1] else None, "min_value": float(r[2]) if r[2] else None, "max_value": float(r[3]) if r[3] else None, "unit": r[4]} for r in cur.fetchall()]
            return jsonify({"success": True, "black_mix": result}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/black-mix/<int:mix_id>/adn", methods=["GET"])
def get_black_mix_adn(mix_id):
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, black_mix_id, adn_text, version, created_at FROM public.black_mix_adn WHERE black_mix_id = %s ORDER BY version DESC LIMIT 1", (mix_id,))
            row = cur.fetchone()
            if not row:
                return jsonify({"success": False, "error": "ADN not found for this Black Mix"}), 404
            adn_id, black_mix_id, adn_text, version, created_at = row
            return jsonify({"success": True, "adn": {"id": adn_id, "black_mix_id": black_mix_id, "version": version, "created_at": created_at.isoformat() if created_at else None, "snapshot": adn_text}}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/black-mix/<int:mix_id>/adn-enriched", methods=["GET"])
def get_black_mix_adn_enriched(mix_id):
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("SELECT id, reference, name, status, document_revision_history FROM black_mixes WHERE id = %s", (mix_id,))
                black_mix = cur.fetchone()
                if not black_mix:
                    return jsonify({"error": "Black mix not found"}), 404
                cur.execute("""
                    SELECT c.matiere_id, c.sub_black_mix_id, c.component_name, c.quantity_value, c.quantity_unit, c.metadata,
                           m.reference, m.nom_matiere, m.type_matiere,
                           bm.reference AS sub_bm_ref, bm.name AS sub_bm_name, f.specifications
                    FROM black_mix_components c
                    LEFT JOIN matieres m ON m.matiere_id = c.matiere_id
                    LEFT JOIN black_mixes bm ON bm.id = c.sub_black_mix_id
                    LEFT JOIN fiches_adn_matieres f ON f.matiere_id = c.matiere_id
                    WHERE c.black_mix_id = %s
                """, (mix_id,))
                components = cur.fetchall()
                for c in components:
                    if not c["specifications"]:
                        c["specifications"] = "Information non disponible"
                cur.execute("SELECT id, step_order, step_name, machine_name, parameters FROM black_mix_process_steps WHERE black_mix_id = %s ORDER BY step_order", (mix_id,))
                process_steps = cur.fetchall()
                cur.execute("""
                    SELECT sm.process_step_id, COALESCE(bm.reference, m.reference) AS reference, COALESCE(bm.name, m.nom_matiere) AS nom_matiere
                    FROM black_mix_step_materials sm
                    JOIN black_mix_process_steps ps ON ps.id = sm.process_step_id
                    LEFT JOIN matieres m ON m.matiere_id = sm.matiere_id
                    LEFT JOIN black_mixes bm ON bm.id = sm.sub_black_mix_id
                    WHERE ps.black_mix_id = %s
                """, (mix_id,))
                step_materials = cur.fetchall()
                materials_by_step = {}
                for row in step_materials:
                    materials_by_step.setdefault(row["process_step_id"], []).append({"reference": row["reference"], "nom_matiere": row["nom_matiere"]})
                for step in process_steps:
                    step["materials"] = materials_by_step.get(step["id"], [])
                cur.execute("SELECT parameter_name, target_value, min_value, max_value, unit FROM black_mix_control_plan WHERE black_mix_id = %s", (mix_id,))
                control_plan = cur.fetchall()
                data_for_ai = {"black_mix_identity": {"reference": black_mix["reference"], "name": black_mix["name"], "status": black_mix["status"], "revision_history": black_mix["document_revision_history"]}, "components": components, "process_steps": process_steps, "step_materials": step_materials, "control_plan": control_plan}
                prompt = f"""Tu es un expert en formulation industrielle de matériaux carbone et graphite.
Génère un "Rapport Technique BLACK MIX ADN" en suivant STRICTEMENT cette structure:
#### 1. Introduction
#### 2. Identité et Vue d'ensemble du Black Mix
#### 3. Architecture et Processus du Black Mix
#### 4. ADN Détaillé des Composants
#### 5. Synthèse de l'Identité Structurelle
RÈGLES: Aucune hallucination. Langue: Français. Style professionnel.
### DONNÉES SOURCE (JSON):
{json.dumps(data_for_ai, indent=2, ensure_ascii=False)}"""
                ai_response = call_groq_with_retry(messages=[{"role": "system", "content": "Tu es un expert en formulation industrielle."}, {"role": "user", "content": prompt}], model="llama-3.3-70b-versatile", temperature=0.2, max_tokens=6000)
                return jsonify({"black_mix": black_mix, "source_data": data_for_ai, "ai_analysis": ai_response.choices[0].message.content if ai_response.choices else ""}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/black-mix/<int:mix_id>/adn-combined", methods=["GET"])
def get_black_mix_adn_combined(mix_id):
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, reference, name, status, created_at, document_revision_history FROM public.black_mixes WHERE id = %s", (mix_id,))
            bm = cur.fetchone()
            if not bm:
                return jsonify({"success": False, "error": "Black Mix not found"}), 404
            black_mix_info = {"id": bm[0], "product_reference": bm[1], "mix_name": bm[2], "status": bm[3], "created_at": bm[4].isoformat() if bm[4] else None, "document_revision_history": bm[5]}
            cur.execute("SELECT adn_text, version, created_at FROM public.black_mix_adn WHERE black_mix_id = %s ORDER BY version DESC LIMIT 1", (mix_id,))
            adn_row = cur.fetchone()
            if not adn_row:
                return jsonify({"success": False, "error": "ADN not found for this Black Mix"}), 404
            base_adn = adn_row[0]
            adn_version = adn_row[1]
            adn_created = adn_row[2].isoformat() if adn_row[2] else None
            components_combined = []
            for comp in base_adn.get("composition", []):
                ref = comp.get("reference")
                entry = {"reference": ref, "material_name": comp.get("material_name", comp.get("component_name", "")), "quantity": comp.get("quantity"), "unit": comp.get("unit"), "metadata": comp.get("metadata", {}), "is_sub_black_mix": comp.get("is_sub_black_mix", False), "adn_matiere": None}
                if ref and not comp.get("is_sub_black_mix"):
                    cur.execute("SELECT fiche_adn_id, nom_matiere, type_matiere, specifications, num_specifications FROM public.fiches_adn_matieres WHERE reference = %s LIMIT 1", (ref,))
                    adn_row2 = cur.fetchone()
                    if adn_row2:
                        entry["adn_matiere"] = {"fiche_adn_id": adn_row2[0], "nom_matiere": adn_row2[1], "type_matiere": adn_row2[2], "specifications": adn_row2[3], "num_specifications": adn_row2[4]}
                if comp.get("is_sub_black_mix") and comp.get("sub_black_mix_adn"):
                    entry["sub_black_mix_adn"] = comp["sub_black_mix_adn"]
                components_combined.append(entry)
            process_steps_combined = []
            for step in base_adn.get("process_steps", []):
                mat_refs = step.get("materials", [])
                materials_detail = []
                for mref in mat_refs:
                    if not mref:
                        continue
                    mat_info = {"reference": mref}
                    for cc in components_combined:
                        if cc["reference"] == mref:
                            mat_info.update({"material_name": cc["material_name"], "quantity": cc["quantity"], "unit": cc["unit"], "is_sub_black_mix": cc["is_sub_black_mix"]})
                            break
                    materials_detail.append(mat_info)
                process_steps_combined.append({"step_order": step.get("step_order"), "step_name": step.get("step_name"), "machine": step.get("machine"), "parameters": step.get("parameters"), "material_references": mat_refs, "materials_detail": materials_detail})
            return jsonify({"success": True, "black_mix": black_mix_info, "adn_version": adn_version, "adn_created_at": adn_created, "composition": components_combined, "composition_flat": base_adn.get("composition_flat", []), "process_steps": process_steps_combined, "control_plan": base_adn.get("control_plan", [])}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


# =============================================================================
# NUANCES — CORE HELPERS
# =============================================================================

def get_matiere_full_adn(cur, matiere_id, reference=None):
    if not matiere_id:
        return None
    cur.execute("""
        SELECT fiche_adn_id, nom_matiere, material_name, reference, type_matiere,
               specifications, num_specifications, date_creation, derniere_modification
        FROM public.fiches_adn_matieres WHERE matiere_id = %s LIMIT 1
    """, (matiere_id,))
    row = cur.fetchone()
    fiche_adn = dict(row) if row else None
    cur.execute("SELECT fiche_id FROM public.fiches_matieres WHERE matiere_id = %s ORDER BY fiche_id DESC", (matiere_id,))
    fiche_rows = cur.fetchall()
    specs = []
    for fr in fiche_rows:
        fiche_id = fr["fiche_id"] if isinstance(fr, dict) else fr[0]
        cur.execute("SELECT spec_id, fiche_id, source_type, donnees, date_creation, derniere_modification FROM public.specifications WHERE fiche_id = %s ORDER BY spec_id", (fiche_id,))
        specs.extend([dict(s) for s in cur.fetchall()])
    cur.execute("""
        SELECT men.id, men.matiere_image_id, men.note_json, men.created_at, mi.image_path
        FROM public.matiere_expert_notes men
        INNER JOIN public.matiere_images mi ON mi.id = men.matiere_image_id
        WHERE mi.matiere_id = %s ORDER BY men.created_at DESC
    """, (matiere_id,))
    expert_notes = [dict(r) for r in cur.fetchall()]
    return {"matiere_id": matiere_id, "fiche_adn": fiche_adn, "specifications": specs, "expert_notes": expert_notes, "summary": {"num_specifications": len(specs), "num_expert_notes": len(expert_notes)}}


def resolve_nuance_ref_lookup(cur, components):
    ref_lookup = {}
    validation_errors = []
    for component in components:
        ref = component.get("reference")
        if not ref:
            validation_errors.append("Component missing reference")
            continue
        if ref in ref_lookup:
            continue
        cur.execute("SELECT id FROM public.nuances WHERE reference = %s LIMIT 1", (ref,))
        row = cur.fetchone()
        if row:
            ref_lookup[ref] = {"type": "nuance", "id": row["id"]}
            continue
        cur.execute("SELECT id FROM public.black_mixes WHERE reference = %s LIMIT 1", (ref,))
        row = cur.fetchone()
        if row:
            ref_lookup[ref] = {"type": "black_mix", "id": row["id"]}
            continue
        cur.execute("SELECT matiere_id FROM public.matieres WHERE reference = %s LIMIT 1", (ref,))
        row = cur.fetchone()
        if row:
            ref_lookup[ref] = {"type": "matiere", "id": row["matiere_id"]}
            continue
        validation_errors.append(f"Reference '{ref}' not found in nuances, black_mixes or matieres")
    return ref_lookup, validation_errors


# =============================================================================
# CUISSON PROGRAMS — HELPERS (UPDATED)
# =============================================================================

def parse_warne_nachbehandlung(raw_value: str):
    """
    Parse '101 25' → (program_number='101', h2_percent=25)
    Formats number as zero-padded string: '1' → '001', '101' → '101'
    K-type kept as-is: 'K000' → 'K000'
    """
    if not raw_value or not raw_value.strip():
        return None, None
    parts = raw_value.strip().split()
    try:
        raw_num = parts[0]
        h2_percent = int(parts[1]) if len(parts) >= 2 else None
        # Format number: numeric → zero-padded 3 digits, K-type → as-is
        if raw_num.upper().startswith('K'):
            program_number = raw_num.upper()
        else:
            program_number = str(int(raw_num)).zfill(3)
        return program_number, h2_percent
    except (ValueError, IndexError):
        return None, None


def get_cuisson_program_by_number(cur, program_number: str):
    """
    Returns full cuisson program from DB including kontrolle column.
    program_number is VARCHAR: '001', '101', 'K000'
    """
    if not program_number:
        return None
    cur.execute(
        """
        SELECT id, program_number, max_temperature, kontrolle, type,
               start_temp, oven_1, oven_2, oven_3, oven_4, oven_5,
               oven_6, oven_7, oven_8, oven_9, oven_10, oven_11,
               oven_12, oven_13, phases_json
        FROM public.cuisson_programs
        WHERE program_number = %s LIMIT 1
        """,
        (program_number,)
    )
    row = cur.fetchone()
    if not row:
        return None
    # Build ovens dict (only non-null)
    ovens = {}
    for i in range(13):
        val = row[6 + i]
        if val:
            ovens[f"oven_{i+1}"] = val
    return {
        "id":                row[0],
        "program_number":    row[1],
        "max_temperature_c": float(row[2]) if row[2] else None,
        "kontrolle":         row[3],
        "type":              row[4],
        "start_temp_c":      float(row[5]) if row[5] else 20,
        "ovens":             ovens,
        "phases":            row[19],
    }


def build_nuance_adn_snapshot(cur, nuance_id, product_reference, nuance_name, _visited=None):
    """
    Recursively builds a fully enriched ADN snapshot for a nuance.
    Includes: cuisson program (with kontrolle), components, mishkarte, control plan, images.
    """
    if _visited is None:
        _visited = set()
    if nuance_id in _visited:
        return {"nuance_id": nuance_id, "product_reference": product_reference, "nuance_name": nuance_name, "error": "Circular reference detected — snapshot truncated here"}
    _visited.add(nuance_id)

    # ── Revision history ─────────────────────────────────────────────────────
    cur.execute("SELECT document_revision_history FROM public.nuances WHERE id = %s", (nuance_id,))
    row = cur.fetchone()
    revision_history = row["document_revision_history"] if row and row["document_revision_history"] else None

    # ── Cuisson program (Wärme-Nachbehandlung) — UPDATED with kontrolle ──────
    cur.execute(
        """
        SELECT n.cuisson_raw, n.cuisson_program_number, n.cuisson_h2_percent,
               cp.type, cp.kontrolle, cp.max_temperature, cp.start_temp, cp.phases_json
        FROM public.nuances n
        LEFT JOIN public.cuisson_programs cp ON cp.id = n.cuisson_program_id
        WHERE n.id = %s
        """,
        (nuance_id,)
    )
    cuisson_row = cur.fetchone()
    cuisson_info = None
    if cuisson_row and cuisson_row["cuisson_raw"]:
        h2 = cuisson_row["cuisson_h2_percent"]
        n2 = (100 - h2) if h2 is not None else None
        cuisson_info = {
            "raw":            cuisson_row["cuisson_raw"],
            "program_number": cuisson_row["cuisson_program_number"],
            "h2_percent":     h2,
            "n2_percent":     n2,
            "atmosphere":     f"H2 {h2}% + N2 {n2}%" if h2 is not None else None,
            "program": {
                "type":              cuisson_row["type"],
                "kontrolle":         cuisson_row["kontrolle"],
                "max_temperature_c": float(cuisson_row["max_temperature"]) if cuisson_row["max_temperature"] else None,
                "start_temp_c":      float(cuisson_row["start_temp"]) if cuisson_row["start_temp"] else 20,
                "phases":            cuisson_row["phases_json"],
            } if cuisson_row["type"] else None,
        }

    # ── Components ───────────────────────────────────────────────────────────
    cur.execute(
        """
        SELECT c.id, c.component_name, c.quantity_value, c.quantity_unit, c.metadata,
               c.matiere_id, c.sub_black_mix_id, c.sub_nuance_id,
               m.reference AS matiere_reference, m.nom_matiere AS matiere_name,
               bm.reference AS sub_bm_reference, bm.name AS sub_bm_name,
               sn.reference AS sub_nuance_reference, sn.name AS sub_nuance_name
        FROM public.nuance_components c
        LEFT JOIN public.matieres    m  ON m.matiere_id = c.matiere_id
        LEFT JOIN public.black_mixes bm ON bm.id        = c.sub_black_mix_id
        LEFT JOIN public.nuances     sn ON sn.id        = c.sub_nuance_id
        WHERE c.nuance_id = %s ORDER BY c.id
        """,
        (nuance_id,)
    )
    raw_components = cur.fetchall()
    components = []
    for r in raw_components:
        comp_id = r["id"]
        component_name = r["component_name"]
        qty = r["quantity_value"]
        unit = r["quantity_unit"]
        metadata = r["metadata"]
        matiere_id = r["matiere_id"]
        sub_bm_id = r["sub_black_mix_id"]
        sub_nuance_id = r["sub_nuance_id"]
        matiere_ref = r["matiere_reference"]
        matiere_name = r["matiere_name"]
        sub_bm_ref = r["sub_bm_reference"]
        sub_bm_name = r["sub_bm_name"]
        sub_nuance_ref = r["sub_nuance_reference"]
        sub_nuance_name = r["sub_nuance_name"]
        if sub_nuance_id is not None:
            comp_type, ref, name = "nuance",    sub_nuance_ref, sub_nuance_name
        elif sub_bm_id is not None:
            comp_type, ref, name = "black_mix", sub_bm_ref,     sub_bm_name
        else:
            comp_type, ref, name = "matiere",   matiere_ref,    matiere_name
        entry = {"id": comp_id, "component_name": component_name, "quantity": float(qty) if qty is not None else None, "unit": unit, "metadata": metadata, "component_type": comp_type, "reference": ref, "material_name": name, "matiere_id": matiere_id, "sub_black_mix_id": sub_bm_id, "sub_nuance_id": sub_nuance_id, "matiere_adn": None, "black_mix_adn": None, "sub_nuance_adn": None}
        if comp_type == "matiere" and matiere_id:
            entry["matiere_adn"] = get_matiere_full_adn(cur, matiere_id, ref)
        elif comp_type == "black_mix" and sub_bm_id:
            entry["black_mix_adn"] = build_black_mix_adn_snapshot(cur, sub_bm_id, sub_bm_ref, sub_bm_name, _visited=set(_visited))
        elif comp_type == "nuance" and sub_nuance_id:
            entry["sub_nuance_adn"] = build_nuance_adn_snapshot(cur, sub_nuance_id, sub_nuance_ref, sub_nuance_name, _visited=set(_visited))
        components.append(entry)

    # ── Process steps (mishkarte) ─────────────────────────────────────────────
    cur.execute("SELECT s.id, s.step_order, s.step_name, s.machine_name, s.parameters FROM public.nuance_process_steps s WHERE s.nuance_id = %s ORDER BY s.step_order", (nuance_id,))
    steps_raw = cur.fetchall()
    process_steps = []
    for s in steps_raw:
        step_id = s["id"]
        step_order = s["step_order"]
        step_name = s["step_name"]
        machine = s["machine_name"]
        parameters = s["parameters"]
        cur.execute("""
            SELECT sm.matiere_id, sm.sub_black_mix_id, sm.sub_nuance_id,
                   m.reference AS matiere_ref, bm.reference AS sub_bm_ref, sn.reference AS sub_nuance_ref
            FROM public.nuance_step_materials sm
            LEFT JOIN public.matieres    m  ON m.matiere_id = sm.matiere_id
            LEFT JOIN public.black_mixes bm ON bm.id        = sm.sub_black_mix_id
            LEFT JOIN public.nuances     sn ON sn.id        = sm.sub_nuance_id
            WHERE sm.process_step_id = %s
        """, (step_id,))
        step_mat_refs = []
        for sm in cur.fetchall():
            ref = sm["sub_nuance_ref"] if sm["sub_nuance_id"] is not None else (sm["sub_bm_ref"] if sm["sub_black_mix_id"] is not None else sm["matiere_ref"])
            if ref:
                step_mat_refs.append(ref)
        process_steps.append({"step_order": step_order, "step_name": step_name, "machine": machine, "parameters": parameters, "materials": step_mat_refs})

    # ── Control plan ─────────────────────────────────────────────────────────
    cur.execute("SELECT parameter_name, target_value, min_value, max_value, unit FROM public.nuance_control_plan WHERE nuance_id = %s ORDER BY parameter_name", (nuance_id,))
    control_plan = [{"parameter_name": r["parameter_name"], "target_value": float(r["target_value"]) if r["target_value"] is not None else None, "min_value": float(r["min_value"]) if r["min_value"] is not None else None, "max_value": float(r["max_value"]) if r["max_value"] is not None else None, "unit": r["unit"]} for r in cur.fetchall()]

    # ── Images + expert notes ─────────────────────────────────────────────────
    cur.execute("""
        SELECT ni.id, ni.image_path, ne.note_json, ne.created_at AS note_created_at
        FROM public.nuance_images ni
        LEFT JOIN public.nuance_expert_notes ne ON ne.nuance_image_id = ni.id
        WHERE ni.nuance_id = %s ORDER BY ni.id
    """, (nuance_id,))
    images = [{"image_id": r["id"], "image_path": r["image_path"], "expert_note": {"note_json": r["note_json"], "created_at": r["note_created_at"].isoformat() if r["note_created_at"] else None} if r["note_json"] else None} for r in cur.fetchall()]

    def flatten(comp_list, depth=0):
        flat = []
        for c in comp_list:
            flat.append({**c, "depth": depth, "matiere_adn": None, "black_mix_adn": None, "sub_nuance_adn": None})
            if c["component_type"] == "nuance" and c["sub_nuance_adn"]:
                flat.extend(flatten(c["sub_nuance_adn"].get("composition", []), depth + 1))
            elif c["component_type"] == "black_mix" and c["black_mix_adn"]:
                for bm_comp in c["black_mix_adn"].get("composition", []):
                    flat.append({**bm_comp, "depth": depth + 1})
        return flat

    return {
        "nuance_id":                 nuance_id,
        "product_reference":         product_reference,
        "nuance_name":               nuance_name,
        "status":                    "draft",
        "document_revision_history": revision_history,
        "created_at":                datetime.now().isoformat(),
        "cuisson":                   cuisson_info,
        "composition":               components,
        "composition_flat":          flatten(components),
        "process_steps":             process_steps,
        "step_materials":            {str(s["step_order"]): s["materials"] for s in process_steps},
        "control_plan":              control_plan,
        "images":                    images,
        "snapshot_timestamp":        datetime.now().isoformat(),
    }


# =============================================================================
# NUANCES — ENDPOINTS
# =============================================================================

@app.route("/nuance/validate-material/<reference>", methods=["GET"])
def validate_nuance_material(reference):
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, name FROM public.nuances WHERE reference = %s LIMIT 1", (reference,))
            row = cur.fetchone()
            if row:
                return jsonify({"reference": reference, "exists": True, "component_type": "nuance", "material_name": row[1], "id": row[0]}), 200
            cur.execute("SELECT id, name FROM public.black_mixes WHERE reference = %s LIMIT 1", (reference,))
            row = cur.fetchone()
            if row:
                return jsonify({"reference": reference, "exists": True, "component_type": "black_mix", "material_name": row[1], "id": row[0]}), 200
            cur.execute("SELECT matiere_id, nom_matiere FROM public.matieres WHERE reference = %s LIMIT 1", (reference,))
            row = cur.fetchone()
            if row:
                return jsonify({"reference": reference, "exists": True, "component_type": "matiere", "material_name": row[1], "id": row[0]}), 200
            return jsonify({"reference": reference, "exists": False, "component_type": None, "material_name": None, "id": None}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/nuance/submit", methods=["POST"])
def submit_nuance():
    if not request.is_json:
        return jsonify({"success": False, "error": "Request body must be JSON"}), 400

    data = request.get_json()

    import logging
    logging.info(f"📥 Payload reçu:\n{json.dumps(data, indent=2)}")

    product_reference  = data.get("product_reference")
    nuance_name        = data.get("nuance_name")
    components         = data.get("components", [])
    process_steps      = data.get("process_steps", [])
    step_materials_map = data.get("step_materials", {})
    control_plan       = data.get("control_plan", [])

    document_revision_history = data.get("document_revision_history")

    warne_raw = data.get("warne_nachbehandlung")
    if warne_raw:
        warne_raw = warne_raw.strip()
        if warne_raw == "":
            warne_raw = None

    # 🔴 VALIDATION
    if not product_reference or not nuance_name:
        return jsonify({
            "success": False,
            "error": "product_reference and nuance_name are required"
        }), 400

    if not process_steps:
        return jsonify({
            "success": False,
            "error": "At least one process_step is required"
        }), 400

    conn = psycopg2.connect(DB_DSN)

    try:
        with conn:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:

                # ───────── VALIDATION DES REFERENCES ─────────
                ref_lookup, validation_errors = resolve_nuance_ref_lookup(cur, components)

                if validation_errors:
                    return jsonify({
                        "success": False,
                        "validation_errors": validation_errors
                    }), 400

                # 🔴 Validation step_materials
                all_step_refs = {
                    ref for refs in step_materials_map.values()
                    for ref in refs if ref
                }

                extra_refs = all_step_refs - set(ref_lookup.keys())

                if extra_refs:
                    extra_lookup, extra_errors = resolve_nuance_ref_lookup(
                        cur, [{"reference": r} for r in extra_refs]
                    )

                    if extra_errors:
                        return jsonify({
                            "success": False,
                            "validation_errors": extra_errors
                        }), 400

                    ref_lookup.update(extra_lookup)

                logging.info(f"✅ Ref lookup: {ref_lookup}")

                # ───────── CUISSON ─────────
                cuisson_program_number = None
                cuisson_h2_percent     = None
                cuisson_program_id     = None

                if warne_raw:
                    cuisson_program_number, cuisson_h2_percent = parse_warne_nachbehandlung(warne_raw)

                    if cuisson_program_number:
                        cur.execute(
                            "SELECT id FROM public.cuisson_programs WHERE program_number = %s LIMIT 1",
                            (cuisson_program_number,)
                        )
                        prog_row = cur.fetchone()

                        if prog_row:
                            cuisson_program_id = prog_row["id"]

                # ───────── INSERT NUANCE ─────────
                cur.execute("""
                    INSERT INTO public.nuances
                        (reference, name, status, created_at,
                         document_revision_history,
                         cuisson_raw, cuisson_program_number,
                         cuisson_h2_percent, cuisson_program_id)
                    VALUES (%s, %s, 'draft', NOW(), %s, %s, %s, %s, %s)
                    RETURNING id
                """, (
                    product_reference,
                    nuance_name,
                    Json(document_revision_history) if document_revision_history else None,
                    warne_raw,
                    cuisson_program_number,
                    cuisson_h2_percent,
                    cuisson_program_id
                ))

                nuance_id = cur.fetchone()["id"]

                # ───────── COMPONENTS ─────────
                for component in components:
                    ref = component.get("reference")
                    resolved = ref_lookup[ref]

                    cur.execute("""
                        INSERT INTO public.nuance_components
                            (nuance_id, matiere_id, sub_black_mix_id, sub_nuance_id,
                             component_name, quantity_value, quantity_unit, metadata)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    """, (
                        nuance_id,
                        resolved["id"] if resolved["type"] == "matiere" else None,
                        resolved["id"] if resolved["type"] == "black_mix" else None,
                        resolved["id"] if resolved["type"] == "nuance" else None,
                        component.get("component_name") or ref,
                        component.get("quantity"),
                        component.get("unit", "kg"),
                        Json(component.get("metadata") or {})
                    ))

                # ───────── PROCESS STEPS ─────────
                for step in process_steps:
                    step_order = step.get("step_order")

                    cur.execute("""
                        INSERT INTO public.nuance_process_steps
                            (nuance_id, step_order, step_name, machine_name, parameters)
                        VALUES (%s, %s, %s, %s, %s)
                        RETURNING id
                    """, (
                        nuance_id,
                        step_order,
                        step.get("step_name"),
                        step.get("machine"),
                        Json(step.get("parameters") or {})
                    ))

                    process_step_id = cur.fetchone()["id"]

                    refs_for_step = step_materials_map.get(str(step_order), [])

                    for ref in refs_for_step:
                        resolved = ref_lookup.get(ref)

                        if not resolved:
                            raise ValueError(f"Reference '{ref}' not resolved")

                        cur.execute("""
                            INSERT INTO public.nuance_step_materials
                                (process_step_id, matiere_id, sub_black_mix_id, sub_nuance_id, created_at)
                            VALUES (%s, %s, %s, %s, NOW())
                        """, (
                            process_step_id,
                            resolved["id"] if resolved["type"] == "matiere" else None,
                            resolved["id"] if resolved["type"] == "black_mix" else None,
                            resolved["id"] if resolved["type"] == "nuance" else None
                        ))

                # ───────── CONTROL PLAN ─────────
                for param in control_plan:
                    cur.execute("""
                        INSERT INTO public.nuance_control_plan
                            (nuance_id, parameter_name, target_value, min_value, max_value, unit)
                        VALUES (%s, %s, %s, %s, %s, %s)
                    """, (
                        nuance_id,
                        param.get("parameter_name"),
                        param.get("target_value"),
                        param.get("min_value"),
                        param.get("max_value"),
                        param.get("unit")
                    ))

                # ───────── ADN ─────────
                adn_snapshot = build_nuance_adn_snapshot(
                    cur, nuance_id, product_reference, nuance_name
                )
                adn_snapshot = serialize_to_json_compatible(adn_snapshot)
                adn_json_str = json.dumps(adn_snapshot, default=str)

                cur.execute("""
                    INSERT INTO public.nuance_adn
                        (nuance_id, adn_text, version, created_at)
                    VALUES (%s, %s::jsonb, 1, NOW())
                    RETURNING id
                """, (nuance_id, adn_json_str))

                adn_id = cur.fetchone()["id"]

                return jsonify({
                    "success": True,
                    "nuance_id": nuance_id,
                    "adn": {
                        "id": adn_id,
                        "version": 1
                    }
                }), 200

    except Exception as e:
        conn.rollback()
        import traceback
        tb = traceback.format_exc()
        logging.error("🔥 ERROR submit_nuance", exc_info=True)

        return jsonify({
            "success": False,
            "error": str(e),
            "type": type(e).__name__,
            "traceback": tb
        }), 500

    finally:
        conn.close()
@app.route("/nuance/list", methods=["GET"])
def list_nuances():
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, reference, name, status, created_at FROM public.nuances ORDER BY created_at DESC")
            rows = cur.fetchall()
            return jsonify({"success": True, "nuances": [{"id": r[0], "product_reference": r[1], "nuance_name": r[2], "status": r[3], "created_at": r[4].isoformat() if r[4] else None} for r in rows]}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/nuance/<int:nuance_id>", methods=["GET"])
def get_nuance_details(nuance_id):
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, reference, name, status, created_at, document_revision_history FROM public.nuances WHERE id = %s", (nuance_id,))
            row = cur.fetchone()
            if not row:
                return jsonify({"success": False, "error": "Nuance not found"}), 404
            result = {"id": row[0], "product_reference": row[1], "nuance_name": row[2], "status": row[3], "created_at": row[4].isoformat() if row[4] else None, "document_revision_history": row[5]}
            snapshot = build_nuance_adn_snapshot(cur, nuance_id, result["product_reference"], result["nuance_name"])
            result["cuisson"]          = snapshot["cuisson"]
            result["components"]       = snapshot["composition"]
            result["components_flat"]  = snapshot["composition_flat"]
            result["process_steps"]    = snapshot["process_steps"]
            result["control_plan"]     = snapshot["control_plan"]
            result["images"]           = [{**img, "image_url": build_image_url(img["image_path"])} if img.get("image_path") else img for img in snapshot.get("images", [])]
            return jsonify({"success": True, "nuance": result}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/nuance/<int:nuance_id>/adn", methods=["GET"])
def get_nuance_adn(nuance_id):
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, nuance_id, adn_text, version, created_at FROM public.nuance_adn WHERE nuance_id = %s ORDER BY version DESC LIMIT 1", (nuance_id,))
            row = cur.fetchone()
            if not row:
                return jsonify({"success": False, "error": "ADN not found for this Nuance"}), 404
            return jsonify({"success": True, "adn": {"id": row[0], "nuance_id": row[1], "version": row[3], "created_at": row[4].isoformat() if row[4] else None, "snapshot": row[2]}}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/nuance/<int:nuance_id>/adn-enriched", methods=["GET"])
def get_nuance_adn_enriched(nuance_id):
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("SELECT id, reference, name, status, document_revision_history FROM public.nuances WHERE id = %s", (nuance_id,))
                nuance = cur.fetchone()
                if not nuance:
                    return jsonify({"error": "Nuance not found"}), 404
                snapshot = build_nuance_adn_snapshot(cur, nuance_id, nuance["reference"], nuance["name"])
                data_for_ai = {"nuance_identity": {"reference": nuance["reference"], "name": nuance["name"], "status": nuance["status"], "revision_history": nuance["document_revision_history"], "cuisson": snapshot.get("cuisson")}, "components": snapshot["composition"], "process_steps": snapshot["process_steps"], "control_plan": snapshot["control_plan"]}
                prompt = f"""Tu es un expert en formulation industrielle de matériaux carbone et graphite.
Génère un "Rapport Technique NUANCE ADN" en suivant STRICTEMENT cette structure:
#### 1. Introduction
#### 2. Identité et Vue d'ensemble de la Nuance
*   **2.1. Informations Générales** : tableau (Référence, Nom, Statut).
*   **2.2. Programme de Cuisson** : Wärme-Nachbehandlung, type, température max, atmosphère H2/N2, phases.
*   **2.3. Historique des Révisions**.
#### 3. Architecture et Processus de la Nuance
*   **3.1. Composition Structurelle** : tableau (Référence, Matériau, Type, Quantité, Rôle).
*   **3.2. Gamme de Fabrication (Mishkarte)** : tableau (Étape, Opération, Machine, Paramètres, Matières).
*   **3.3. Analyse d'impact du processus**.
*   **3.4. Plan de Contrôle**.
#### 4. ADN Détaillé des Composants
Pour CHAQUE composant : Référence, Type, Fonction, Spécifications complètes.
#### 5. Synthèse de l'Identité Structurelle
RÈGLES: Aucune hallucination. Langue: Français. Style professionnel.
### DONNÉES SOURCE (JSON):
{json.dumps(data_for_ai, indent=2, ensure_ascii=False)}"""
                ai_response = call_groq_with_retry(messages=[{"role": "system", "content": "Tu es un expert en formulation industrielle."}, {"role": "user", "content": prompt}], model="llama-3.3-70b-versatile", temperature=0.2, max_tokens=6000)
                return jsonify({"nuance": nuance, "source_data": data_for_ai, "ai_analysis": ai_response.choices[0].message.content if ai_response.choices else ""}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/nuance/<int:nuance_id>/adn-combined", methods=["GET"])
def get_nuance_adn_combined(nuance_id):
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, reference, name, status, created_at, document_revision_history FROM public.nuances WHERE id = %s", (nuance_id,))
            row = cur.fetchone()
            if not row:
                return jsonify({"success": False, "error": "Nuance not found"}), 404
            nuance_info = {"id": row[0], "product_reference": row[1], "nuance_name": row[2], "status": row[3], "created_at": row[4].isoformat() if row[4] else None, "document_revision_history": row[5]}
            snapshot = build_nuance_adn_snapshot(cur, nuance_id, nuance_info["product_reference"], nuance_info["nuance_name"])
            for img in snapshot.get("images", []):
                if img.get("image_path"):
                    img["image_url"] = build_image_url(img["image_path"])
            return jsonify({"success": True, "nuance": nuance_info, "snapshot_version": "live", "cuisson": snapshot["cuisson"], "composition": snapshot["composition"], "composition_flat": snapshot["composition_flat"], "process_steps": snapshot["process_steps"], "control_plan": snapshot["control_plan"], "images": snapshot["images"]}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


# =============================================================================
# NUANCE IMAGES — UPLOAD + SEARCH PAR SIMILARITÉ
# =============================================================================

def search_similar_nuances_in_db(query_embedding: np.ndarray, top_k: int = 5):
    query_vec = query_embedding.tolist()
    sql = """
        SELECT ni.id, ni.image_path, ni.nuance_id, n.name AS nuance_name, n.reference,
               (1 - (ni.embedding <=> %s)) AS similarity
        FROM public.nuance_images ni
        JOIN public.nuances n ON n.id = ni.nuance_id
        ORDER BY ni.embedding <=> %s LIMIT %s;
    """
    conn = get_db_conn()
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute(sql, (query_vec, query_vec, top_k))
            return [dict(r) for r in cur.fetchall()]
    finally:
        conn.close()


@app.route("/nuance/upload-image", methods=["POST"])
def upload_nuance_image():
    nuance_id = request.form.get("nuance_id")
    if not nuance_id:
        return jsonify({"success": False, "error": "nuance_id is required"}), 400
    if "image" not in request.files:
        return jsonify({"success": False, "error": "No image file provided"}), 400
    file = request.files["image"]
    if not file or not allowed_file(file.filename):
        return jsonify({"success": False, "error": "Invalid file type. Allowed: png, jpg, jpeg"}), 400
    expert_note_raw = request.form.get("expert_note")
    expert_note = None
    if expert_note_raw:
        try:
            expert_note = json.loads(expert_note_raw)
        except Exception:
            return jsonify({"success": False, "error": "expert_note must be valid JSON"}), 400
    conn = None
    try:
        filename_safe = secure_filename(file.filename)
        unique_filename = f"nuance_{nuance_id}_{uuid.uuid4().hex}_{filename_safe}"
        file_path = IMAGES_DIR / unique_filename
        file_bytes = file.read()
        with open(file_path, "wb") as f:
            f.write(file_bytes)
        img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        embedding = compute_embedding_from_pil(img)
        conn = get_db_conn()
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT name, reference FROM public.nuances WHERE id = %s", (int(nuance_id),))
            nuance = cur.fetchone()
            if not nuance:
                return jsonify({"success": False, "error": "Nuance not found"}), 404
            cur.execute("INSERT INTO public.nuance_images (nuance_id, image_path, embedding, nuance_name, reference, created_at) VALUES (%s, %s, %s, %s, %s, NOW()) RETURNING id",
                        (int(nuance_id), str(file_path), embedding.tolist(), nuance["name"], nuance["reference"]))
            image_id = cur.fetchone()["id"]
            note_id = None
            if expert_note:
                cur.execute("INSERT INTO public.nuance_expert_notes (nuance_image_id, note_json, created_at) VALUES (%s, %s, NOW()) RETURNING id", (image_id, Json(expert_note)))
                note_id = cur.fetchone()["id"]
            conn.commit()
        return jsonify({"success": True, "message": "Image uploaded and embedding computed successfully", "image_id": image_id, "image_url": build_image_url(str(file_path)), "nuance_id": int(nuance_id), "note_id": note_id}), 200
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        if conn:
            conn.close()


@app.route("/nuance/search-similar", methods=["POST"])
def search_similar_nuances():
    data = request.get_json(silent=True) or {}
    top_k = int(data.get("top_k", 5))
    if top_k < 1 or top_k > 50:
        return jsonify({"success": False, "error": "invalid_top_k"}), 400
    temp_filename = data.get("temp_filename")
    file_id = data.get("file_id")
    download_link = data.get("download_link")
    provided = [bool(download_link), bool(temp_filename), bool(file_id)]
    if sum(provided) != 1:
        return jsonify({"success": False, "error": "Provide exactly ONE of: download_link, temp_filename, file_id"}), 400
    img = None
    if download_link:
        try:
            r = requests.get(download_link, timeout=20)
            r.raise_for_status()
            img = Image.open(io.BytesIO(r.content)).convert("RGB")
        except Exception as e:
            return jsonify({"success": False, "error": "download_link_failed", "message": str(e)}), 400
    elif temp_filename:
        file_path = TEMP_UPLOAD_DIR / temp_filename
        if not file_path.exists():
            return jsonify({"success": False, "error": "temp_file_not_found"}), 404
        img = Image.open(file_path).convert("RGB")
    elif file_id:
        if not client:
            return jsonify({"success": False, "error": "openai_not_configured"}), 400
        img = Image.open(io.BytesIO(client.files.content(file_id).read())).convert("RGB")
    try:
        query_embedding = compute_embedding_from_pil(img)
        rows = search_similar_nuances_in_db(query_embedding, top_k=top_k)
        results = [{"id": r["id"], "image_url": build_image_url(r["image_path"]), "nuance_id": r["nuance_id"], "nuance_name": r["nuance_name"], "reference": r["reference"], "similarity": float(r["similarity"]) if r["similarity"] is not None else None} for r in rows]
        return jsonify({"success": True, "results": results}), 200
    except Exception as e:
        return jsonify({"success": False, "error": "search_failed", "message": str(e)}), 500


# =============================================================================
# CUISSON PROGRAMS — ENDPOINTS (UPDATED with kontrolle + VARCHAR)
# =============================================================================

@app.route("/cuisson-programs", methods=["GET"])
def list_cuisson_programs():
    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT id, program_number, max_temperature, kontrolle, type,
                       start_temp, oven_1, oven_2, oven_3, oven_4, oven_5,
                       oven_6, oven_7, oven_8, oven_9, oven_10, oven_11,
                       oven_12, oven_13, phases_json
                FROM public.cuisson_programs
                ORDER BY program_number
            """)
            rows = [dict(r) for r in cur.fetchall()]
            return jsonify({"success": True, "count": len(rows), "programs": rows}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/cuisson-programs/<string:program_number>", methods=["GET"])
def get_cuisson_program_detail(program_number):
    """
    Get cuisson program by number.
    Accepts: '101', '001', 'K000' — auto-formats numeric to zero-padded.
    """
    conn = psycopg2.connect(DB_DSN)
    try:
        # Normalize: numeric → zero-padded, K-type → uppercase as-is
        if program_number.upper().startswith('K'):
            formatted = program_number.upper()
        else:
            try:
                formatted = str(int(program_number)).zfill(3)
            except ValueError:
                formatted = program_number

        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            program = get_cuisson_program_by_number(cur, formatted)
            if not program:
                return jsonify({"success": False, "error": f"Programme '{formatted}' non trouvé"}), 404
            return jsonify({"success": True, "program": program}), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/cuisson-programs/parse", methods=["GET"])
def parse_cuisson_field():
    """
    Parse Wärme-Nachbehandlung field and return full program details.
    Example: GET /cuisson-programs/parse?value=101 25
    """
    raw_value = request.args.get("value", "").strip()
    if not raw_value:
        return jsonify({"success": False, "error": "Paramètre 'value' requis. Ex: ?value=101 25"}), 400

    program_number, h2_percent = parse_warne_nachbehandlung(raw_value)
    n2_percent = (100 - h2_percent) if h2_percent is not None else None

    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            program = get_cuisson_program_by_number(cur, program_number)
        return jsonify({
            "success":        True,
            "raw_value":      raw_value,
            "program_number": program_number,
            "h2_percent":     h2_percent,
            "n2_percent":     n2_percent,
            "atmosphere":     f"H2 {h2_percent}% + N2 {n2_percent}%" if h2_percent is not None else None,
            "program_found":  program is not None,
            "program":        program,
        }), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


@app.route("/nuance/<int:nuance_id>/set-cuisson", methods=["POST"])
def set_nuance_cuisson(nuance_id):
    """
    Associate a cuisson program to a nuance.
    Body: { "warne_nachbehandlung": "101 25" }
    """
    data = request.get_json() or {}
    raw_value = data.get("warne_nachbehandlung", "").strip()
    if not raw_value:
        return jsonify({"success": False, "error": "Champ 'warne_nachbehandlung' requis. Ex: '101 25'"}), 400

    program_number, h2_percent = parse_warne_nachbehandlung(raw_value)
    if not program_number:
        return jsonify({"success": False, "error": f"Format invalide: '{raw_value}'. Attendu: '101 25'"}), 400

    conn = psycopg2.connect(DB_DSN)
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT id, reference FROM public.nuances WHERE id = %s", (nuance_id,))
            nuance = cur.fetchone()
            if not nuance:
                return jsonify({"success": False, "error": "Nuance non trouvée"}), 404

            cur.execute("SELECT id FROM public.cuisson_programs WHERE program_number = %s LIMIT 1", (program_number,))
            prog_row = cur.fetchone()
            cuisson_program_id = prog_row["id"] if prog_row else None

            cur.execute(
                """
                UPDATE public.nuances
                SET cuisson_raw            = %s,
                    cuisson_program_number = %s,
                    cuisson_h2_percent     = %s,
                    cuisson_program_id     = %s,
                    updated_at             = NOW()
                WHERE id = %s
                """,
                (raw_value, program_number, h2_percent, cuisson_program_id, nuance_id)
            )
            conn.commit()

            program = get_cuisson_program_by_number(cur, program_number) if cuisson_program_id else None
            n2_percent = (100 - h2_percent) if h2_percent is not None else None

            return jsonify({
                "success":             True,
                "message":             f"Programme de cuisson '{raw_value}' associé à la nuance {nuance['reference']}",
                "nuance_id":           nuance_id,
                "cuisson_raw":         raw_value,
                "program_number":      program_number,
                "h2_percent":          h2_percent,
                "n2_percent":          n2_percent,
                "atmosphere":          f"H2 {h2_percent}% + N2 {n2_percent}%",
                "program_found_in_db": cuisson_program_id is not None,
                "program":             program,
            }), 200
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()

@app.route("/nuance/<int:nuance_id>/update", methods=["PUT"])
def update_nuance(nuance_id):
    if not request.is_json:
        return jsonify({"success": False, "error": "Request body must be JSON"}), 400
    data = request.get_json()

    conn = psycopg2.connect(DB_DSN)
    try:
        with conn:
            with conn.cursor() as cur:

                # ── 0. Vérifier existence ─────────────────────────────────────
                cur.execute("SELECT id, reference FROM public.nuances WHERE id = %s", (nuance_id,))
                existing = cur.fetchone()
                if not existing:
                    return jsonify({"success": False, "error": "Nuance not found"}), 404

                # ── 1. Champs identité ────────────────────────────────────────
                product_reference         = data.get("product_reference")
                nuance_name               = data.get("nuance_name")
                document_revision_history = data.get("document_revision_history")
                warne_raw                 = (data.get("warne_nachbehandlung") or "").strip()

                if not product_reference or not nuance_name:
                    return jsonify({"success": False, "error": "product_reference and nuance_name are required"}), 400

                # ── 2. Parser le programme de cuisson ────────────────────────
                cuisson_program_number = cuisson_h2_percent = cuisson_program_id = None
                if warne_raw:
                    cuisson_program_number, cuisson_h2_percent = parse_warne_nachbehandlung(warne_raw)
                    if cuisson_program_number:
                        cur.execute(
                            "SELECT id FROM public.cuisson_programs WHERE program_number = %s LIMIT 1",
                            (cuisson_program_number,)
                        )
                        prog_row = cur.fetchone()
                        if prog_row:
                            cuisson_program_id = prog_row[0]

                # ── 3. Mettre à jour la table nuances ────────────────────────
                cur.execute(
                    """
                    UPDATE public.nuances
                    SET reference                  = %s,
                        name                       = %s,
                        document_revision_history  = %s,
                        cuisson_raw                = %s,
                        cuisson_program_number     = %s,
                        cuisson_h2_percent         = %s,
                        cuisson_program_id         = %s,
                        updated_at                 = NOW()
                    WHERE id = %s
                    """,
                    (
                        product_reference,
                        nuance_name,
                        Json(document_revision_history) if document_revision_history else None,
                        warne_raw or None,
                        cuisson_program_number,
                        cuisson_h2_percent,
                        cuisson_program_id,
                        nuance_id
                    )
                )

                # ── 4. Résoudre et valider toutes les références ──────────────
                components         = data.get("components", [])
                process_steps      = data.get("process_steps", [])
                step_materials_map = data.get("step_materials", {})
                control_plan       = data.get("control_plan", [])

                ref_lookup, validation_errors = resolve_nuance_ref_lookup(cur, components)
                if validation_errors:
                    return jsonify({"success": False, "validation_errors": validation_errors}), 400

                all_step_refs = {ref for refs in step_materials_map.values() for ref in refs if ref}
                extra_refs    = all_step_refs - set(ref_lookup.keys())
                if extra_refs:
                    extra_lookup, extra_errors = resolve_nuance_ref_lookup(
                        cur, [{"reference": r} for r in extra_refs]
                    )
                    if extra_errors:
                        return jsonify({"success": False, "validation_errors": extra_errors}), 400
                    ref_lookup.update(extra_lookup)

                # ── 5. Supprimer l'ancienne composition / étapes / contrôle ──
                # Récupérer les IDs des anciennes étapes pour supprimer leurs matières
                cur.execute(
                    "SELECT id FROM public.nuance_process_steps WHERE nuance_id = %s",
                    (nuance_id,)
                )
                old_step_ids = [r[0] for r in cur.fetchall()]
                if old_step_ids:
                    cur.execute(
                        "DELETE FROM public.nuance_step_materials WHERE process_step_id = ANY(%s)",
                        (old_step_ids,)
                    )
                cur.execute("DELETE FROM public.nuance_process_steps  WHERE nuance_id = %s", (nuance_id,))
                cur.execute("DELETE FROM public.nuance_components      WHERE nuance_id = %s", (nuance_id,))
                cur.execute("DELETE FROM public.nuance_control_plan    WHERE nuance_id = %s", (nuance_id,))

                # ── 6. Réinsérer les composants ───────────────────────────────
                for component in components:
                    ref      = component.get("reference")
                    resolved = ref_lookup[ref]
                    cur.execute(
                        """
                        INSERT INTO public.nuance_components
                            (nuance_id, matiere_id, sub_black_mix_id, sub_nuance_id,
                             component_name, quantity_value, quantity_unit, metadata)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        """,
                        (
                            nuance_id,
                            resolved["id"] if resolved["type"] == "matiere"    else None,
                            resolved["id"] if resolved["type"] == "black_mix"  else None,
                            resolved["id"] if resolved["type"] == "nuance"     else None,
                            component.get("component_name") or ref,
                            component.get("quantity"),
                            component.get("unit", "kg"),
                            Json(component.get("metadata", {}))
                        )
                    )

                # ── 7. Réinsérer les étapes + step_materials ──────────────────
                if not process_steps:
                    return jsonify({"success": False, "error": "At least one process_step is required"}), 400

                for step in process_steps:
                    step_order = step.get("step_order")
                    cur.execute(
                        """
                        INSERT INTO public.nuance_process_steps
                            (nuance_id, step_order, step_name, machine_name, parameters)
                        VALUES (%s, %s, %s, %s, %s) RETURNING id
                        """,
                        (
                            nuance_id,
                            step_order,
                            step.get("step_name"),
                            step.get("machine"),
                            Json(step.get("parameters", {}))
                        )
                    )
                    process_step_id = cur.fetchone()[0]

                    refs_for_step = step_materials_map.get(str(step_order), [])
                    if not refs_for_step:
                        raise ValueError(
                            f"Step '{step.get('step_name')}' (order {step_order}) "
                            f"has no materials in step_materials"
                        )

                    seen_ids = set()
                    for ref in refs_for_step:
                        if not ref:
                            continue
                        resolved = ref_lookup.get(ref)
                        if not resolved:
                            raise ValueError(f"Reference '{ref}' in step_materials not resolved")
                        dedup_key = (resolved["type"], resolved["id"])
                        if dedup_key in seen_ids:
                            continue
                        seen_ids.add(dedup_key)
                        cur.execute(
                            """
                            INSERT INTO public.nuance_step_materials
                                (process_step_id, matiere_id, sub_black_mix_id, sub_nuance_id, created_at)
                            VALUES (%s, %s, %s, %s, NOW())
                            """,
                            (
                                process_step_id,
                                resolved["id"] if resolved["type"] == "matiere"   else None,
                                resolved["id"] if resolved["type"] == "black_mix" else None,
                                resolved["id"] if resolved["type"] == "nuance"    else None,
                            )
                        )

                # ── 8. Réinsérer le plan de contrôle ─────────────────────────
                for param in control_plan:
                    cur.execute(
                        """
                        INSERT INTO public.nuance_control_plan
                            (nuance_id, parameter_name, target_value, min_value, max_value, unit)
                        VALUES (%s, %s, %s, %s, %s, %s)
                        """,
                        (
                            nuance_id,
                            param.get("parameter_name"),
                            param.get("target_value"),
                            param.get("min_value"),
                            param.get("max_value"),
                            param.get("unit")
                        )
                    )

                # ── 9. Reconstruire et versionner l'ADN ──────────────────────
                new_snapshot = build_nuance_adn_snapshot(
                    cur, nuance_id, product_reference, nuance_name
                )

                cur.execute(
                    "SELECT version FROM public.nuance_adn WHERE nuance_id = %s ORDER BY version DESC LIMIT 1",
                    (nuance_id,)
                )
                last_version_row = cur.fetchone()
                next_version     = (last_version_row[0] + 1) if last_version_row else 1

                cur.execute(
                    """
                    INSERT INTO public.nuance_adn
                        (nuance_id, adn_text, version, created_at)
                    VALUES (%s, %s, %s, NOW()) RETURNING id
                    """,
                    (nuance_id, Json(new_snapshot), next_version)
                )
                new_adn_id = cur.fetchone()[0]

                return jsonify({
                    "success":           True,
                    "message":           f"Nuance '{nuance_name}' updated successfully",
                    "nuance_id":         nuance_id,
                    "product_reference": product_reference,
                    "component_types":   {ref: info["type"] for ref, info in ref_lookup.items()},
                    "adn": {
                        "id":      new_adn_id,
                        "version": next_version
                    },
                    "cuisson": {
                        "raw":           warne_raw or None,
                        "program_number": cuisson_program_number,
                        "h2_percent":    cuisson_h2_percent,
                        "program_found": cuisson_program_id is not None
                    }
                }), 200

    except ValueError as ve:
        conn.rollback()
        return jsonify({"success": False, "error": str(ve)}), 400
    except Exception as e:
        conn.rollback()
        logging.error(f"Update Nuance error: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()

# =============================================================================
# XLS → XLSX CONVERSION + SHEET EXTRACTION TO JSON
# =============================================================================

def _xlrd_cell_to_str(cell, book):
    """Convert an xlrd cell to a clean string value."""
    import xlrd
    value = cell.value
    if cell.ctype == xlrd.XL_CELL_DATE:
        try:
            dt_tuple = xlrd.xldate_as_tuple(value, book.datemode)
            return datetime(*dt_tuple).strftime("%Y-%m-%d")
        except Exception:
            return str(value)
    elif cell.ctype == xlrd.XL_CELL_NUMBER:
        if value == int(value):
            return str(int(value))
        return str(value)
    elif cell.ctype == xlrd.XL_CELL_BOOLEAN:
        return "TRUE" if value else "FALSE"
    elif cell.ctype == xlrd.XL_CELL_EMPTY:
        return ""
    return str(value).strip()


def _extract_sheet_raw(xls_sheet, book):
    """Extract all non-empty cells from an xlrd sheet."""
    cells = []
    all_text_parts = []
    rows_compact = []
    for row_idx in range(xls_sheet.nrows):
        row_dict = {}
        for col_idx in range(xls_sheet.ncols):
            value = _xlrd_cell_to_str(xls_sheet.cell(row_idx, col_idx), book)
            if value != "":
                cells.append({"row": row_idx + 1, "col": col_idx + 1, "value": value})
                all_text_parts.append(value)
                row_dict[f"col_{col_idx + 1}"] = value
        if row_dict:
            row_dict["_row"] = row_idx + 1
            rows_compact.append(row_dict)
    return cells, all_text_parts, rows_compact


def _parse_sollwerte(xls_sheet, book):
    """Parse the Sollwerte sheet into structured specification data."""
    import xlrd
    result = {
        "vormischung": "",
        "reference": "",
        "revisions": []
    }
    # Row 0: Vormischung / reference
    if xls_sheet.nrows > 0:
        if xls_sheet.ncols > 0:
            result["vormischung"] = _xlrd_cell_to_str(xls_sheet.cell(0, 0), book)
        if xls_sheet.ncols > 1:
            result["reference"] = _xlrd_cell_to_str(xls_sheet.cell(0, 1), book)

    # Parse revision blocks. A new block starts at "geändert am/von", "erstellt am",
    # or "Änderungsgrund". The topmost block in the sheet is the latest revision.
    current_revision = None
    i = 1
    while i < xls_sheet.nrows:
        col0 = _xlrd_cell_to_str(xls_sheet.cell(i, 0), book).lower()

        if "geändert am" in col0 or "erstellt am" in col0:
            if current_revision:
                result["revisions"].append(current_revision)
            date_val = _xlrd_cell_to_str(xls_sheet.cell(i, 1), book) if xls_sheet.ncols > 1 else ""
            author = _xlrd_cell_to_str(xls_sheet.cell(i, 2), book) if xls_sheet.ncols > 2 else ""
            current_revision = {
                "date": date_val,
                "author": author,
                "change_reason": "",
                "specifications": []
            }
            i += 1
            continue

        if "änderungsgrund" in col0:
            # Each Änderungsgrund starts a new spec block
            if current_revision and current_revision["specifications"]:
                result["revisions"].append(current_revision)
            reason = _xlrd_cell_to_str(xls_sheet.cell(i, 1), book) if xls_sheet.ncols > 1 else ""
            if not current_revision:
                current_revision = {"date": "", "author": "", "change_reason": reason, "specifications": []}
            else:
                current_revision = {
                    "date": current_revision.get("date", ""),
                    "author": current_revision.get("author", ""),
                    "change_reason": reason,
                    "specifications": []
                }
            i += 1
            continue

        # Check for min/max header row
        col1 = _xlrd_cell_to_str(xls_sheet.cell(i, 1), book).lower() if xls_sheet.ncols > 1 else ""
        if col1 == "min":
            i += 1
            continue

        # Spec rows (e.g. "Schüttdichte [g/L]", "> 630 µm [%]", etc.)
        if col0 and ("µm" in col0 or "schüttdichte" in col0 or "dichte" in col0):
            name = _xlrd_cell_to_str(xls_sheet.cell(i, 0), book)
            min_val = _xlrd_cell_to_str(xls_sheet.cell(i, 1), book) if xls_sheet.ncols > 1 else ""
            max_val = _xlrd_cell_to_str(xls_sheet.cell(i, 2), book) if xls_sheet.ncols > 2 else ""
            spec = {"parameter": name, "min": min_val, "max": max_val}
            if current_revision:
                current_revision["specifications"].append(spec)
            i += 1
            continue

        i += 1

    if current_revision:
        result["revisions"].append(current_revision)

    # Active specs = first revision block that has specs with actual min/max values
    result["active_specifications"] = []
    for rev in result["revisions"]:
        specs = [s for s in rev.get("specifications", []) if s.get("min") or s.get("max")]
        if specs:
            result["active_specifications"] = specs
            break

    return result


def _parse_data_sheet(xls_sheet, book):
    """Parse the reference/data sheet (e.g. '049 91') into structured measurement records."""
    import xlrd
    result = {
        "type": "",
        "reference": "",
        "method": "",
        "headers": [],
        "measurements": []
    }

    if xls_sheet.nrows < 7:
        return result

    # Row 0: VM/HM, reference, method
    result["type"] = _xlrd_cell_to_str(xls_sheet.cell(0, 0), book)
    result["reference"] = _xlrd_cell_to_str(xls_sheet.cell(0, 1), book)
    if xls_sheet.ncols > 4:
        result["method"] = _xlrd_cell_to_str(xls_sheet.cell(0, 4), book)

    # Data rows start at row 7 (index 7) based on structure:
    # Row 3: top headers (Datum, Charge, Prüfer, Schüttdichte, Siebe)
    # Row 4: sub-headers (FREI, sieve labels)
    # Row 5: IST/SOLL labels
    # Row 6: min/max labels
    # Row 7+: data
    for row_idx in range(7, xls_sheet.nrows):
        row_num = _xlrd_cell_to_str(xls_sheet.cell(row_idx, 0), book)
        if not row_num:
            continue

        datum = _xlrd_cell_to_str(xls_sheet.cell(row_idx, 1), book)
        charge = _xlrd_cell_to_str(xls_sheet.cell(row_idx, 2), book)
        pruefer = _xlrd_cell_to_str(xls_sheet.cell(row_idx, 3), book)

        record = {
            "nr": row_num,
            "datum": datum,
            "charge": charge,
            "pruefer": pruefer,
            "schuettdichte": _xlrd_cell_to_str(xls_sheet.cell(row_idx, 4), book),
        }

        # Sieve IST values only (SOLL is already in Sollwerte)
        if xls_sheet.ncols > 10:
            record["sieve_630"] = _xlrd_cell_to_str(xls_sheet.cell(row_idx, 10), book)
        if xls_sheet.ncols > 13:
            record["sieve_355"] = _xlrd_cell_to_str(xls_sheet.cell(row_idx, 13), book)
        if xls_sheet.ncols > 16:
            record["sieve_90"] = _xlrd_cell_to_str(xls_sheet.cell(row_idx, 16), book)
        if xls_sheet.ncols > 19:
            record["sieve_lt90"] = _xlrd_cell_to_str(xls_sheet.cell(row_idx, 19), book)

        # Remove empty values to save space
        record = {k: v for k, v in record.items() if v}

        result["measurements"].append(record)

    return result


def _compute_measurement_stats(measurements, active_specs=None):
    """Compute summary statistics for measurement data to avoid sending all rows."""
    if not measurements:
        return {"count": 0}

    # Numeric fields to aggregate
    fields = ["schuettdichte", "sieve_630", "sieve_355", "sieve_90", "sieve_lt90"]
    stats = {"count": len(measurements)}

    for field in fields:
        values = []
        for m in measurements:
            raw = m.get(field, "")
            if raw:
                try:
                    values.append(float(str(raw).replace(",", ".")))
                except (ValueError, TypeError):
                    pass
        if values:
            avg = sum(values) / len(values)
            stats[field] = {
                "count": len(values),
                "min": round(min(values), 2),
                "max": round(max(values), 2),
                "mean": round(avg, 2),
                "std": round((sum((x - avg) ** 2 for x in values) / len(values)) ** 0.5, 2)
            }
            # Count out-of-spec if specs are available
            if active_specs:
                spec = None
                spec_map = {
                    "schuettdichte": "schüttdichte",
                    "sieve_630": "630",
                    "sieve_355": "355",
                    "sieve_90": "> 90",
                    "sieve_lt90": "< 90"
                }
                search_key = spec_map.get(field, "")
                for s in active_specs:
                    if search_key in s.get("parameter", "").lower():
                        spec = s
                        break
                if spec:
                    spec_min = None
                    spec_max = None
                    try:
                        if spec.get("min"):
                            spec_min = float(str(spec["min"]).replace(",", "."))
                    except (ValueError, TypeError):
                        pass
                    try:
                        if spec.get("max"):
                            spec_max = float(str(spec["max"]).replace(",", "."))
                    except (ValueError, TypeError):
                        pass
                    oos = 0
                    for v in values:
                        if (spec_min is not None and v < spec_min) or (spec_max is not None and v > spec_max):
                            oos += 1
                    stats[field]["out_of_spec"] = oos

    # Date range
    dates = [m.get("datum", "") for m in measurements if m.get("datum")]
    if dates:
        stats["date_range"] = {"first": dates[0], "last": dates[-1]}

    return stats


def _find_non_conformities(measurements, active_specs):
    """Find charges (lots) whose measurement values exceed min/max from Sollwerte.
    A charge may appear multiple times if it was re-tested after a non-conformity."""
    if not measurements or not active_specs:
        return []

    # Build spec lookup: field_name -> (min, max)
    spec_map = {
        "schuettdichte": "schüttdichte",
        "sieve_630": "630",
        "sieve_355": "355",
        "sieve_90": "> 90",
        "sieve_lt90": "< 90"
    }
    spec_limits = {}
    for field, search_key in spec_map.items():
        for s in active_specs:
            if search_key in s.get("parameter", "").lower():
                spec_min = None
                spec_max = None
                try:
                    if s.get("min"):
                        spec_min = float(str(s["min"]).replace(",", "."))
                except (ValueError, TypeError):
                    pass
                try:
                    if s.get("max"):
                        spec_max = float(str(s["max"]).replace(",", "."))
                except (ValueError, TypeError):
                    pass
                if spec_min is not None or spec_max is not None:
                    spec_limits[field] = (spec_min, spec_max, s.get("parameter", ""))
                break

    non_conf = []
    for m in measurements:
        charge = m.get("charge", "")
        if not charge:
            continue
        failures = []
        for field, (s_min, s_max, param_name) in spec_limits.items():
            raw = m.get(field, "")
            if not raw:
                continue
            try:
                val = float(str(raw).replace(",", "."))
            except (ValueError, TypeError):
                continue
            if (s_min is not None and val < s_min) or (s_max is not None and val > s_max):
                failures.append({
                    "parameter": param_name,
                    "value": val,
                    "min": s_min,
                    "max": s_max
                })
        if failures:
            non_conf.append({
                "charge": charge,
                "nr": m.get("nr", ""),
                "datum": m.get("datum", ""),
                "failures": failures
            })

    return non_conf


@app.route("/convert_xls_to_json", methods=["POST"])
def convert_xls_to_json():
    """
    Accepts an .xls file via JSON (OpenAI GPT integration), parses and extracts content.

    For VM/HM files (filename starts with VM or HM): parses the 'Sollwerte' sheet
    (specifications) and the reference data sheet (next to Sollwerte) into structured JSON.

    For other files: extracts all cells as raw data.

    JSON body:
        openaiFileIdRefs – list of file refs with download_link, name, id (required)
        sheet_name       – name of the sheet to extract (optional, default first sheet)
        mode             – 'summary' (default, stats + last N rows) or 'full' (all rows)
        max_rows         – max measurement rows to include (default 5 in summary mode)
    """
    import xlrd

    data = request.get_json(silent=True) or {}
    refs = data.get("openaiFileIdRefs")
    sheet_name = data.get("sheet_name", None)
    mode = str(data.get("mode", "summary")).lower()
    max_rows = int(data.get("max_rows", 5))

    if not refs or not isinstance(refs, list) or len(refs) == 0:
        return jsonify({"success": False, "error": "Missing openaiFileIdRefs. Send a JSON body with openaiFileIdRefs list."}), 400

    file_ref = refs[0]
    download_link = file_ref.get("download_link")
    original_name = file_ref.get("name") or "uploaded_file.xls"
    file_id = file_ref.get("id")

    if not original_name.lower().endswith(".xls"):
        return jsonify({"success": False, "error": "Only .xls files are accepted."}), 400

    # Download the file
    xls_bytes = None
    if download_link:
        try:
            r = requests.get(download_link, timeout=30)
            r.raise_for_status()
            xls_bytes = r.content
        except Exception as e:
            return jsonify({"success": False, "error": f"download_link failed: {e}"}), 400

    if xls_bytes is None and file_id and client:
        try:
            xls_bytes = client.files.content(file_id).read()
        except Exception as e:
            return jsonify({"success": False, "error": f"OpenAI file download failed: {e}"}), 400

    if xls_bytes is None:
        return jsonify({"success": False, "error": "Could not download file. Provide a valid download_link or file id."}), 400

    original_name = secure_filename(original_name)

    try:
        xls_book = xlrd.open_workbook(file_contents=xls_bytes, formatting_info=False)
        all_sheets = xls_book.sheet_names()

        # --- Detect VM/HM files ---
        is_vm_hm = original_name.upper().startswith("VM") or original_name.upper().startswith("HM")
        has_sollwerte = "Sollwerte" in all_sheets

        if is_vm_hm and has_sollwerte:
            # Structured parsing for VM/HM files
            sollwerte_sheet = xls_book.sheet_by_name("Sollwerte")
            sollwerte_data = _parse_sollwerte(sollwerte_sheet, xls_book)

            # The data sheet is the one right after Sollwerte
            sollwerte_idx = all_sheets.index("Sollwerte")
            data_sheet_name = all_sheets[sollwerte_idx + 1] if sollwerte_idx + 1 < len(all_sheets) else None

            data_sheet_result = None
            if data_sheet_name:
                data_sheet = xls_book.sheet_by_name(data_sheet_name)
                data_sheet_result = _parse_data_sheet(data_sheet, xls_book)



            # Build data_sheet response based on mode
            ds_response = None
            non_conformities = []
            if data_sheet_name and data_sheet_result:
                all_measurements = data_sheet_result["measurements"]
                active_specs = sollwerte_data.get("active_specifications", [])

                # Find charges with values outside spec
                non_conformities = _find_non_conformities(all_measurements, active_specs)

                ds_response = {
                    "sheet_name": data_sheet_name,
                    "type": data_sheet_result["type"],
                    "reference": data_sheet_result["reference"],
                    "method": data_sheet_result["method"],
                    "total_measurements": len(all_measurements),
                }
                if mode == "full":
                    ds_response["measurements"] = all_measurements
                else:
                    # Summary mode: stats + last N rows
                    ds_response["statistics"] = _compute_measurement_stats(all_measurements, active_specs)
                    ds_response["last_measurements"] = all_measurements[-max_rows:]

            return jsonify({
                "success": True,
                "file_type": "VM/HM",
                "source_file": original_name,
                "available_sheets": all_sheets,
                "mode": mode,
                "sollwerte": sollwerte_data,
                "non_conformities": non_conformities,
                "data_sheet": ds_response
            }), 200

        # --- Generic parsing for non-VM/HM files ---
        if sheet_name:
            if sheet_name not in all_sheets:
                return jsonify({
                    "success": False,
                    "error": f"Sheet '{sheet_name}' not found. Available sheets: {all_sheets}"
                }), 400
            xls_sheet = xls_book.sheet_by_name(sheet_name)
        else:
            xls_sheet = xls_book.sheet_by_index(0)
            sheet_name = xls_sheet.name

        cells, all_text_parts, rows_compact = _extract_sheet_raw(xls_sheet, xls_book)

        return jsonify({
            "success": True,
            "file_type": "generic",
            "source_file": original_name,
            "sheet": sheet_name,
            "available_sheets": all_sheets,
            "rows": rows_compact
        }), 200

    except xlrd.biffh.XLRDError as e:
        return jsonify({"success": False, "error": f"Invalid or corrupt .xls file: {e}"}), 400
    except Exception as e:
        logging.error(f"convert_xls_to_json error: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500

# =============================================================================
# MAIN
# =============================================================================
if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--host", type=str, default="0.0.0.0")
    parser.add_argument("--port", type=int, default=int(os.getenv("PORT", "5000")))
    parser.add_argument("--debug", action="store_true")
    args = parser.parse_args()
    cleanup_thread = Thread(target=cleanup_old_files, daemon=True)
    cleanup_thread.start()
    print("🧹 Background cleanup task started (runs every 30 minutes)")
    app.run(host=args.host, port=args.port, debug=args.debug)
